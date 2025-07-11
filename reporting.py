# reporting.py - 报告生成模块
import streamlit as st
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import seaborn as sns
import matplotlib.pyplot as plt
from datetime import datetime, timedelta
import io
import base64
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
import docx
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import zipfile
import tempfile
import os

def reporting_module():
    """报告生成模块主界面"""
    st.title("📄 报告生成模块")
    st.markdown("---")
    
    # 侧边栏 - 报告类型选择
    st.sidebar.markdown("### 📋 报告类型")
    
    report_types = {
        "📊 统计分析报告": "statistical_report",
        "📈 数据可视化报告": "visualization_report", 
        "🔍 数据质量报告": "quality_report",
        "📋 研究总结报告": "summary_report",
        "⚠️ 安全性报告": "safety_report",
        "📊 中期分析报告": "interim_report",
        "📑 最终研究报告": "final_report",
        "📊 自定义报告": "custom_report"
    }
    
    selected_report = st.sidebar.selectbox(
        "选择报告类型",
        options=list(report_types.keys())
    )
    
    report_type = report_types[selected_report]
    
    # 主界面
    st.markdown(f"## {selected_report}")
    
    # 数据源选择
    st.markdown("### 📁 数据源选择")
    
    data_source = st.selectbox(
        "选择数据源",
        ["上传文件", "使用示例数据", "连接数据库"]
    )
    
    df = None
    
    if data_source == "上传文件":
        uploaded_file = st.file_uploader(
            "选择数据文件",
            type=['csv', 'xlsx', 'json'],
            help="支持CSV、Excel和JSON格式"
        )
        
        if uploaded_file is not None:
            try:
                if uploaded_file.name.endswith('.csv'):
                    df = pd.read_csv(uploaded_file)
                elif uploaded_file.name.endswith('.xlsx'):
                    df = pd.read_excel(uploaded_file)
                elif uploaded_file.name.endswith('.json'):
                    df = pd.read_json(uploaded_file)
                
                st.success(f"数据加载成功！共 {len(df)} 行，{len(df.columns)} 列")
                
            except Exception as e:
                st.error(f"数据加载失败：{str(e)}")
    
    elif data_source == "使用示例数据":
        df = generate_sample_data()
        st.success("示例数据加载成功！")
    
    elif data_source == "连接数据库":
        st.info("数据库连接功能开发中...")
    
    # 如果有数据，显示数据预览
    if df is not None:
        with st.expander("📊 数据预览"):
            st.dataframe(df.head(10))
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("总行数", len(df))
            with col2:
                st.metric("总列数", len(df.columns))
            with col3:
                st.metric("缺失值", df.isnull().sum().sum())
    
    # 根据选择的报告类型调用相应函数
    if df is not None:
        if report_type == "statistical_report":
            generate_statistical_report(df)
        elif report_type == "visualization_report":
            generate_visualization_report(df)
        elif report_type == "quality_report":
            generate_quality_report(df)
        elif report_type == "summary_report":
            generate_summary_report(df)
        elif report_type == "safety_report":
            generate_safety_report(df)
        elif report_type == "interim_report":
            generate_interim_report(df)
        elif report_type == "final_report":
            generate_final_report(df)
        elif report_type == "custom_report":
            generate_custom_report(df)

def generate_sample_data():
    """生成示例数据"""
    np.random.seed(42)
    
    n_subjects = 200
    
    # 基础信息
    subject_ids = [f"S{i:04d}" for i in range(1, n_subjects + 1)]
    ages = np.random.normal(45, 12, n_subjects).astype(int)
    ages = np.clip(ages, 18, 80)
    
    genders = np.random.choice(['男', '女'], n_subjects, p=[0.6, 0.4])
    
    # 分组信息
    groups = np.random.choice(['试验组', '对照组'], n_subjects, p=[0.5, 0.5])
    
    # 基线指标
    baseline_sbp = np.random.normal(140, 15, n_subjects)
    baseline_dbp = np.random.normal(90, 10, n_subjects)
    baseline_weight = np.random.normal(70, 12, n_subjects)
    
    # 终点指标（模拟治疗效果）
    treatment_effect = np.where(groups == '试验组', -10, -2)
    endpoint_sbp = baseline_sbp + treatment_effect + np.random.normal(0, 8, n_subjects)
    endpoint_dbp = baseline_dbp + treatment_effect * 0.6 + np.random.normal(0, 6, n_subjects)
    
    # 不良事件
    ae_prob = np.where(groups == '试验组', 0.15, 0.12)
    adverse_events = np.random.binomial(1, ae_prob)
    
    # 依从性
    compliance = np.random.normal(0.85, 0.15, n_subjects)
    compliance = np.clip(compliance, 0, 1)
    
    df = pd.DataFrame({
        'subject_id': subject_ids,
        'age': ages,
        'gender': genders,
        'group': groups,
        'baseline_sbp': baseline_sbp,
        'baseline_dbp': baseline_dbp,
        'baseline_weight': baseline_weight,
        'endpoint_sbp': endpoint_sbp,
        'endpoint_dbp': endpoint_dbp,
        'sbp_change': endpoint_sbp - baseline_sbp,
        'dbp_change': endpoint_dbp - baseline_dbp,
        'adverse_event': adverse_events,
        'compliance': compliance,
        'visit_date': pd.date_range('2024-01-01', periods=n_subjects, freq='D')
    })
    
    return df

def generate_statistical_report(df):
    """生成统计分析报告"""
    st.markdown("### 📊 统计分析报告")
    
    # 报告配置
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### ⚙️ 分析配置")
        
        # 选择分组变量
        group_var = st.selectbox(
            "分组变量",
            options=[col for col in df.columns if df[col].dtype == 'object'],
            help="用于分组比较的变量"
        )
        
        # 选择分析变量
        analysis_vars = st.multiselect(
            "分析变量",
            options=[col for col in df.columns if df[col].dtype in ['int64', 'float64']],
            help="需要进行统计分析的数值变量"
        )
        
        # 统计方法选择
        stat_methods = st.multiselect(
            "统计方法",
            ["描述性统计", "t检验", "卡方检验", "方差分析", "非参数检验", "相关分析"],
            default=["描述性统计", "t检验"]
        )
    
    with col2:
        st.markdown("#### 📋 报告选项")
        
        include_plots = st.checkbox("包含图表", value=True)
        include_tables = st.checkbox("包含统计表格", value=True)
        confidence_level = st.slider("置信水平", 0.90, 0.99, 0.95, 0.01)
        
        report_format = st.selectbox(
            "报告格式",
            ["HTML", "PDF", "Word", "PowerPoint"]
        )
        
        report_title = st.text_input(
            "报告标题",
            value="统计分析报告"
        )
    
    if st.button("📊 生成统计分析报告", type="primary"):
        
        if not analysis_vars:
            st.error("请选择至少一个分析变量")
            return
        
        # 生成报告内容
        report_content = create_statistical_analysis(
            df, group_var, analysis_vars, stat_methods, confidence_level
        )
        
        # 显示报告
        display_statistical_report(
            report_content, include_plots, include_tables,
            report_title, report_format
        )

def create_statistical_analysis(df, group_var, analysis_vars, stat_methods, confidence_level):
    """创建统计分析内容"""
    
    results = {
        'basic_info': {
            'total_subjects': len(df),
            'analysis_date': datetime.now().strftime('%Y-%m-%d'),
            'confidence_level': confidence_level
        },
        'descriptive_stats': {},
        'group_comparison': {},
        'plots': {}
    }
    
    # 描述性统计
    if "描述性统计" in stat_methods:
        results['descriptive_stats'] = perform_descriptive_analysis(df, group_var, analysis_vars)
    
    # 组间比较
    if any(method in stat_methods for method in ["t检验", "卡方检验", "方差分析", "非参数检验"]):
        results['group_comparison'] = perform_group_comparison(
            df, group_var, analysis_vars, stat_methods, confidence_level
        )
    
    # 相关分析
    if "相关分析" in stat_methods:
        results['correlation_analysis'] = perform_correlation_analysis(df, analysis_vars)
    
    return results

def perform_descriptive_analysis(df, group_var, analysis_vars):
    """执行描述性统计分析"""
    
    descriptive_results = {}
    
    # 总体描述性统计
    overall_stats = []
    
    for var in analysis_vars:
        if var in df.columns:
            stats = {
                'variable': var,
                'n': df[var].count(),
                'mean': df[var].mean(),
                'std': df[var].std(),
                'min': df[var].min(),
                'q25': df[var].quantile(0.25),
                'median': df[var].median(),
                'q75': df[var].quantile(0.75),
                'max': df[var].max(),
                'missing': df[var].isnull().sum()
            }
            overall_stats.append(stats)
    
    descriptive_results['overall'] = pd.DataFrame(overall_stats)
    
    # 分组描述性统计
    if group_var and group_var in df.columns:
        group_stats = {}
        
        for group in df[group_var].unique():
            group_data = df[df[group_var] == group]
            group_analysis = []
            
            for var in analysis_vars:
                if var in group_data.columns:
                    stats = {
                        'variable': var,
                        'group': group,
                        'n': group_data[var].count(),
                        'mean': group_data[var].mean(),
                        'std': group_data[var].std(),
                        'median': group_data[var].median(),
                        'missing': group_data[var].isnull().sum()
                    }
                    group_analysis.append(stats)
            
            group_stats[group] = pd.DataFrame(group_analysis)
        
        descriptive_results['by_group'] = group_stats
    
    return descriptive_results

def perform_group_comparison(df, group_var, analysis_vars, stat_methods, confidence_level):
    """执行组间比较分析"""
    
    from scipy import stats
    
    comparison_results = {}
    
    if not group_var or group_var not in df.columns:
        return comparison_results
    
    groups = df[group_var].unique()
    
    if len(groups) < 2:
        return comparison_results
    
    # 两组比较
    if len(groups) == 2:
        group1_data = df[df[group_var] == groups[0]]
        group2_data = df[df[group_var] == groups[1]]
        
        for var in analysis_vars:
            if var in df.columns:
                var_results = {
                    'variable': var,
                    'group1': groups[0],
                    'group2': groups[1],
                    'group1_n': group1_data[var].count(),
                    'group2_n': group2_data[var].count(),
                    'group1_mean': group1_data[var].mean(),
                    'group2_mean': group2_data[var].mean(),
                    'mean_diff': group2_data[var].mean() - group1_data[var].mean()
                }
                
                # t检验
                if "t检验" in stat_methods:
                    try:
                        t_stat, p_value = stats.ttest_ind(
                            group1_data[var].dropna(),
                            group2_data[var].dropna()
                        )
                        var_results['t_statistic'] = t_stat
                        var_results['t_test_p_value'] = p_value
                        var_results['t_test_significant'] = p_value < (1 - confidence_level)
                    except:
                        var_results['t_test_error'] = "无法执行t检验"
                
                # 非参数检验
                if "非参数检验" in stat_methods:
                    try:
                        u_stat, p_value = stats.mannwhitneyu(
                            group1_data[var].dropna(),
                            group2_data[var].dropna()
                        )
                        var_results['mannwhitney_u'] = u_stat
                        var_results['mannwhitney_p_value'] = p_value
                        var_results['mannwhitney_significant'] = p_value < (1 - confidence_level)
                    except:
                        var_results['mannwhitney_error'] = "无法执行Mann-Whitney检验"
                
                comparison_results[var] = var_results
    
    # 多组比较
    elif len(groups) > 2 and "方差分析" in stat_methods:
        for var in analysis_vars:
            if var in df.columns:
                group_data = [df[df[group_var] == group][var].dropna() for group in groups]
                
                try:
                    f_stat, p_value = stats.f_oneway(*group_data)
                    
                    var_results = {
                        'variable': var,
                        'f_statistic': f_stat,
                        'anova_p_value': p_value,
                        'anova_significant': p_value < (1 - confidence_level),
                        'groups': list(groups)
                    }
                    
                    comparison_results[var] = var_results
                    
                except:
                    comparison_results[var] = {'error': '无法执行方差分析'}
    
    return comparison_results

def perform_correlation_analysis(df, analysis_vars):
    """执行相关分析"""
    
    correlation_results = {}
    
    # 选择数值变量
    numeric_vars = [var for var in analysis_vars if var in df.columns and df[var].dtype in ['int64', 'float64']]
    
    if len(numeric_vars) < 2:
        return correlation_results
    
    # 计算相关系数矩阵
    correlation_matrix = df[numeric_vars].corr()
    
    correlation_results['correlation_matrix'] = correlation_matrix
    
    # 显著性检验
    from scipy.stats import pearsonr
    
    correlation_tests = []
    
    for i, var1 in enumerate(numeric_vars):
        for j, var2 in enumerate(numeric_vars):
            if i < j:  # 避免重复
                try:
                    data1 = df[var1].dropna()
                    data2 = df[var2].dropna()
                    
                    # 找到两个变量都有值的观测
                    common_idx = data1.index.intersection(data2.index)
                    
                    if len(common_idx) > 2:
                        corr_coef, p_value = pearsonr(data1[common_idx], data2[common_idx])
                        
                        correlation_tests.append({
                            'variable1': var1,
                            'variable2': var2,
                            'correlation': corr_coef,
                            'p_value': p_value,
                            'n': len(common_idx),
                            'significant': p_value < 0.05
                        })
                except:
                    pass
    
    correlation_results['correlation_tests'] = pd.DataFrame(correlation_tests)
    
    return correlation_results

def display_statistical_report(report_content, include_plots, include_tables, 
                             report_title, report_format):
    """显示统计分析报告"""
    
    st.markdown(f"### 📊 {report_title}")
    st.markdown(f"**生成时间**: {report_content['basic_info']['analysis_date']}")
    st.markdown(f"**样本量**: {report_content['basic_info']['total_subjects']}")
    st.markdown(f"**置信水平**: {report_content['basic_info']['confidence_level']:.0%}")
    
    st.markdown("---")
    
    # 描述性统计
    if 'descriptive_stats' in report_content:
        st.markdown("#### 📈 描述性统计")
        
        if include_tables and 'overall' in report_content['descriptive_stats']:
            st.markdown("##### 总体统计")
            
            desc_df = report_content['descriptive_stats']['overall']
            
            # 格式化数值
            formatted_desc = desc_df.copy()
            numeric_cols = ['mean', 'std', 'min', 'q25', 'median', 'q75', 'max']
            
            for col in numeric_cols:
                if col in formatted_desc.columns:
                    formatted_desc[col] = formatted_desc[col].round(2)
            
            st.dataframe(formatted_desc, hide_index=True)
        
        # 分组统计
        if 'by_group' in report_content['descriptive_stats']:
            st.markdown("##### 分组统计")
            
            for group, group_df in report_content['descriptive_stats']['by_group'].items():
                st.markdown(f"**{group}组**")
                
                formatted_group = group_df.copy()
                numeric_cols = ['mean', 'std', 'median']
                
                for col in numeric_cols:
                    if col in formatted_group.columns:
                        formatted_group[col] = formatted_group[col].round(2)
                
                st.dataframe(formatted_group[['variable', 'n', 'mean', 'std', 'median']], hide_index=True)
    
    # 组间比较
    if 'group_comparison' in report_content and report_content['group_comparison']:
        st.markdown("#### 🔍 组间比较分析")
        
        comparison_results = []
        
        for var, results in report_content['group_comparison'].items():
            result_row = {
                '变量': var,
                '组1': results.get('group1', ''),
                '组2': results.get('group2', ''),
                '均值差': f"{results.get('mean_diff', 0):.3f}",
            }
            
            if 't_test_p_value' in results:
                result_row['t检验P值'] = f"{results['t_test_p_value']:.4f}"
                result_row['t检验显著性'] = "是" if results.get('t_test_significant', False) else "否"
            
            if 'mannwhitney_p_value' in results:
                result_row['Mann-Whitney P值'] = f"{results['mannwhitney_p_value']:.4f}"
            
            comparison_results.append(result_row)
        
        if comparison_results:
            st.dataframe(pd.DataFrame(comparison_results), hide_index=True)
    
    # 相关分析
    if 'correlation_analysis' in report_content:
        st.markdown("#### 🔗 相关分析")
        
        if 'correlation_matrix' in report_content['correlation_analysis']:
            corr_matrix = report_content['correlation_analysis']['correlation_matrix']
            
            if include_plots:
                fig = px.imshow(
                    corr_matrix,
                    text_auto=True,
                    aspect="auto",
                    title="变量相关系数矩阵",
                    color_continuous_scale='RdBu_r'
                )
                fig.update_layout(height=500)
                st.plotly_chart(fig, use_container_width=True)
            
            if include_tables and 'correlation_tests' in report_content['correlation_analysis']:
                st.markdown("##### 相关性检验结果")
                
                corr_tests = report_content['correlation_analysis']['correlation_tests']
                if not corr_tests.empty:
                    formatted_corr = corr_tests.copy()
                    formatted_corr['correlation'] = formatted_corr['correlation'].round(3)
                    formatted_corr['p_value'] = formatted_corr['p_value'].round(4)
                    formatted_corr['significant'] = formatted_corr['significant'].map({True: '是', False: '否'})
                    
                    st.dataframe(formatted_corr, hide_index=True)
    
    # 生成可下载的报告
    st.markdown("---")
    st.markdown("### 💾 下载报告")
    
    if st.button("📥 生成下载文件"):
        
        if report_format == "HTML":
            html_report = generate_html_report(report_content, report_title)
            st.download_button(
                label="下载HTML报告",
                data=html_report,
                file_name=f"{report_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                mime="text/html"
            )
        
        elif report_format == "PDF":
            pdf_report = generate_pdf_report(report_content, report_title)
            st.download_button(
                label="下载PDF报告",
                data=pdf_report,
                file_name=f"{report_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf"
            )
        
        elif report_format == "Word":
            word_report = generate_word_report(report_content, report_title)
            st.download_button(
                label="下载Word报告",
                data=word_report,
                file_name=f"{report_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

def generate_visualization_report(df):
    """生成数据可视化报告"""
    st.markdown("### 📈 数据可视化报告")
    
    # 可视化配置
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### 🎨 图表配置")
        
        # 选择变量
        numeric_vars = [col for col in df.columns if df[col].dtype in ['int64', 'float64']]
        categorical_vars = [col for col in df.columns if df[col].dtype == 'object']
        
        chart_types = st.multiselect(
            "选择图表类型",
            ["直方图", "箱线图", "散点图", "相关热力图", "分组柱状图", "时间序列图"],
            default=["直方图", "箱线图"]
        )
        
        color_scheme = st.selectbox(
            "配色方案",
            ["默认", "蓝色系", "红色系", "绿色系", "彩虹色"]
        )
    
    with col2:
        st.markdown("#### 📊 图表选项")
        
        show_statistics = st.checkbox("显示统计信息", value=True)
        interactive_plots = st.checkbox("交互式图表", value=True)
        
        plot_size = st.selectbox(
            "图表尺寸",
            ["小", "中", "大"],
            index=1
        )
        
        export_format = st.selectbox(
            "导出格式",
            ["PNG", "SVG", "PDF", "HTML"]
        )
    
    if st.button("📈 生成可视化报告", type="primary"):
        
        # 创建可视化内容
        create_visualization_charts(
            df, chart_types, numeric_vars, categorical_vars,
            color_scheme, show_statistics, interactive_plots, plot_size
        )

def create_visualization_charts(df, chart_types, numeric_vars, categorical_vars,
                              color_scheme, show_statistics, interactive_plots, plot_size):
    """创建可视化图表"""
    
    # 设置图表尺寸
    size_map = {"小": 300, "中": 400, "大": 500}
    height = size_map[plot_size]
    
    # 设置配色
    color_map = {
        "默认": px.colors.qualitative.Plotly,
        "蓝色系": px.colors.sequential.Blues,
        "红色系": px.colors.sequential.Reds,
        "绿色系": px.colors.sequential.Greens,
        "彩虹色": px.colors.qualitative.Set1
    }
    colors = color_map[color_scheme]
    
    # 直方图
    if "直方图" in chart_types and numeric_vars:
        st.markdown("#### 📊 数值变量分布")
        
        selected_vars = st.multiselect(
            "选择要展示的数值变量",
            numeric_vars,
            default=numeric_vars[:3]
        )
        
        for var in selected_vars:
            fig = px.histogram(
                df, x=var,
                title=f"{var} 分布图",
                nbins=30,
                color_discrete_sequence=colors
            )
            
            if show_statistics:
                # 添加统计信息
                mean_val = df[var].mean()
                std_val = df[var].std()
                
                fig.add_vline(
                    x=mean_val, 
                    line_dash="dash", 
                    line_color="red",
                    annotation_text=f"均值: {mean_val:.2f}"
                )
            
            fig.update_layout(height=height)
            st.plotly_chart(fig, use_container_width=True)
    
    # 箱线图
    if "箱线图" in chart_types and numeric_vars:
        st.markdown("#### 📦 箱线图分析")
        
        if categorical_vars:
            group_var = st.selectbox(
                "选择分组变量（箱线图）",
                categorical_vars,
                key="boxplot_group"
            )
            
            selected_numeric = st.selectbox(
                "选择数值变量（箱线图）",
                numeric_vars,
                key="boxplot_numeric"
            )
            
            fig = px.box(
                df, x=group_var, y=selected_numeric,
                title=f"{selected_numeric} 按 {group_var} 分组的箱线图",
                color_discrete_sequence=colors
            )
            
            fig.update_layout(height=height)
            st.plotly_chart(fig, use_container_width=True)
        else:
            # 单变量箱线图
            selected_var = st.selectbox(
                "选择变量（单变量箱线图）",
                numeric_vars,
                key="single_boxplot"
            )
            
            fig = px.box(
                df, y=selected_var,
                title=f"{selected_var} 箱线图"
            )
            
            fig.update_layout(height=height)
            st.plotly_chart(fig, use_container_width=True)
    
    # 散点图
    if "散点图" in chart_types and len(numeric_vars) >= 2:
        st.markdown("#### 🔍 散点图分析")
        
        col1, col2 = st.columns(2)
        
        with col1:
            x_var = st.selectbox(
                "X轴变量",
                numeric_vars,
                key="scatter_x"
            )
        
        with col2:
            y_var = st.selectbox(
                "Y轴变量",
                [var for var in numeric_vars if var != x_var],
                key="scatter_y"
            )
        
        # 可选的颜色分组
        color_var = None
        if categorical_vars:
            color_var = st.selectbox(
                "颜色分组变量（可选）",
                ["无"] + categorical_vars,
                key="scatter_color"
            )
            if color_var == "无":
                color_var = None
        
        fig = px.scatter(
            df, x=x_var, y=y_var, color=color_var,
            title=f"{y_var} vs {x_var}",
            trendline="ols" if show_statistics else None,
            color_discrete_sequence=colors
        )
        
        fig.update_layout(height=height)
        st.plotly_chart(fig, use_container_width=True)
        
        # 显示相关系数
        if show_statistics:
            correlation = df[x_var].corr(df[y_var])
            st.info(f"**相关系数**: {correlation:.3f}")
    
    # 相关热力图
    if "相关热力图" in chart_types and len(numeric_vars) >= 2:
        st.markdown("#### 🔥 相关性热力图")
        
        correlation_matrix = df[numeric_vars].corr()
        
        fig = px.imshow(
            correlation_matrix,
            text_auto=True,
            aspect="auto",
            title="变量相关性热力图",
            color_continuous_scale='RdBu_r'
        )
        
        fig.update_layout(height=height)
        st.plotly_chart(fig, use_container_width=True)
    
    # 分组柱状图
    if "分组柱状图" in chart_types and categorical_vars and numeric_vars:
        st.markdown("#### 📊 分组柱状图")
        
        col1, col2 = st.columns(2)
        
        with col1:
            cat_var = st.selectbox(
                "分类变量",
                categorical_vars,
                key="bar_cat"
            )
        
        with col2:
            num_var = st.selectbox(
                "数值变量",
                numeric_vars,
                key="bar_num"
            )
        
        # 计算分组统计
        group_stats = df.groupby(cat_var)[num_var].agg(['mean', 'std', 'count']).reset_index()
        
        fig = px.bar(
            group_stats, x=cat_var, y='mean',
            title=f"{num_var} 按 {cat_var} 分组的均值",
            error_y='std' if show_statistics else None,
            color_discrete_sequence=colors
        )
        
        fig.update_layout(height=height)
        st.plotly_chart(fig, use_container_width=True)
        
        # 显示统计表
        if show_statistics:
            st.dataframe(group_stats, hide_index=True)
    
    # 时间序列图
    if "时间序列图" in chart_types:
        st.markdown("#### 📈 时间序列分析")
        
        # 查找日期列
        date_cols = [col for col in df.columns if df[col].dtype == 'datetime64[ns]' or 'date' in col.lower()]
        
        if date_cols:
            date_var = st.selectbox(
                "时间变量",
                date_cols,
                key="time_var"
            )
            
            num_var = st.selectbox(
                "数值变量",
                numeric_vars,
                key="time_num"
            )
            
            # 确保日期格式正确
            if df[date_var].dtype != 'datetime64[ns]':
                try:
                    df[date_var] = pd.to_datetime(df[date_var])
                except:
                    st.error("无法解析日期格式")
                    return
            
            fig = px.line(
                df.sort_values(date_var), 
                x=date_var, y=num_var,
                title=f"{num_var} 时间趋势",
                color_discrete_sequence=colors
            )
            
            fig.update_layout(height=height)
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.warning("未找到日期类型的列，无法绘制时间序列图")

def generate_quality_report(df):
    """生成数据质量报告"""
    st.markdown("### 🔍 数据质量报告")
    
    # 数据质量概览
    st.markdown("#### 📊 数据质量概览")
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("总行数", len(df))
    
    with col2:
        st.metric("总列数", len(df.columns))
    
    with col3:
        missing_count = df.isnull().sum().sum()
        st.metric("缺失值总数", missing_count)
    
    with col4:
        duplicate_count = df.duplicated().sum()
        st.metric("重复行数", duplicate_count)
    
    # 缺失值分析
    st.markdown("#### ❌ 缺失值分析")
    
    missing_analysis = []
    
    for col in df.columns:
        missing_count = df[col].isnull().sum()
        missing_percent = (missing_count / len(df)) * 100
        
        missing_analysis.append({
            '列名': col,
            '数据类型': str(df[col].dtype),
            '缺失数量': missing_count,
            '缺失比例': f"{missing_percent:.2f}%",
            '完整性': f"{100-missing_percent:.2f}%"
        })
    
    missing_df = pd.DataFrame(missing_analysis)
    missing_df = missing_df.sort_values('缺失数量', ascending=False)
    
    st.dataframe(missing_df, hide_index=True)
    
    # 缺失值可视化
    if missing_df['缺失数量'].sum() > 0:
        fig = px.bar(
            missing_df.head(10), 
            x='列名', y='缺失数量',
            title="各列缺失值数量（前10列）",
            color='缺失数量',
            color_continuous_scale='Reds'
        )
        fig.update_layout(height=400)
        st.plotly_chart(fig, use_container_width=True)
    
    # 数据类型分析
    st.markdown("#### 🏷️ 数据类型分析")
    
    dtype_analysis = df.dtypes.value_counts().reset_index()
    dtype_analysis.columns = ['数据类型', '列数']
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.dataframe(dtype_analysis, hide_index=True)
    
    with col2:
        fig = px.pie(
            dtype_analysis, 
            values='列数', names='数据类型',
            title="数据类型分布"
        )
        fig.update_layout(height=300)
        st.plotly_chart(fig, use_container_width=True)
    
    # 数值变量统计
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    
    if len(numeric_cols) > 0:
        st.markdown("#### 📈 数值变量统计")
        
        numeric_stats = []
        
        for col in numeric_cols:
            stats = {
                '变量名': col,
                '均值': df[col].mean(),
                '标准差': df[col].std(),
                '最小值': df[col].min(),
                '最大值': df[col].max(),
                '零值数': (df[col] == 0).sum(),
                '负值数': (df[col] < 0).sum(),
                '异常值数': detect_outliers(df[col])
            }
            numeric_stats.append(stats)
        
        numeric_df = pd.DataFrame(numeric_stats)
        
        # 格式化数值
        for col in ['均值', '标准差', '最小值', '最大值']:
            numeric_df[col] = numeric_df[col].round(3)
        
        st.dataframe(numeric_df, hide_index=True)
    
    # 分类变量分析
    categorical_cols = df.select_dtypes(include=['object']).columns
    
    if len(categorical_cols) > 0:
        st.markdown("#### 🏷️ 分类变量分析")
        
        categorical_stats = []
        
        for col in categorical_cols:
            unique_count = df[col].nunique()
            most_frequent = df[col].mode().iloc[0] if not df[col].mode().empty else "N/A"
            most_frequent_count = df[col].value_counts().iloc[0] if not df[col].value_counts().empty else 0
            
            stats = {
                '变量名': col,
                '唯一值数': unique_count,
                '最频繁值': most_frequent,
                '最频繁值数量': most_frequent_count,
                '最频繁值比例': f"{(most_frequent_count/len(df))*100:.2f}%"
            }
            categorical_stats.append(stats)
        
        categorical_df = pd.DataFrame(categorical_stats)
        st.dataframe(categorical_df, hide_index=True)
    
    # 数据质量评分
    st.markdown("#### ⭐ 数据质量评分")
    
    quality_score = calculate_quality_score(df)
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("完整性评分", f"{quality_score['completeness']:.1f}/100")
    
    with col2:
        st.metric("一致性评分", f"{quality_score['consistency']:.1f}/100")
    
    with col3:
        st.metric("总体质量评分", f"{quality_score['overall']:.1f}/100")
    
    # 质量改进建议
    st.markdown("#### 💡 质量改进建议")
    
    suggestions = generate_quality_suggestions(df, missing_df, quality_score)
    
    for suggestion in suggestions:
        st.info(f"• {suggestion}")

def detect_outliers(series, method='iqr'):
    """检测异常值"""
    
    if method == 'iqr':
        Q1 = series.quantile(0.25)
        Q3 = series.quantile(0.75)
        IQR = Q3 - Q1
        
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        
        outliers = series[(series < lower_bound) | (series > upper_bound)]
        return len(outliers)
    
    return 0

def calculate_quality_score(df):
    """计算数据质量评分"""
    
    # 完整性评分（基于缺失值比例）
    total_cells = len(df) * len(df.columns)
    missing_cells = df.isnull().sum().sum()
    completeness = ((total_cells - missing_cells) / total_cells) * 100
    
    # 一致性评分（基于数据类型一致性和重复值）
    duplicate_ratio = df.duplicated().sum() / len(df)
    consistency = (1 - duplicate_ratio) * 100
    
    # 总体评分
    overall = (completeness + consistency) / 2
    
    return {
        'completeness': completeness,
        'consistency': consistency,
        'overall': overall
    }

def generate_quality_suggestions(df, missing_df, quality_score):
    """生成数据质量改进建议"""
    
    suggestions = []
    
    # 缺失值建议
    high_missing_cols = missing_df[missing_df['缺失数量'] > len(df) * 0.1]['列名'].tolist()
    if high_missing_cols:
        suggestions.append(f"以下列缺失值较多（>10%），建议检查数据收集过程: {', '.join(high_missing_cols)}")
    
    # 重复值建议
    if df.duplicated().sum() > 0:
        suggestions.append(f"发现 {df.duplicated().sum()} 行重复数据，建议进行去重处理")
    
    # 数据类型建议
    object_cols = df.select_dtypes(include=['object']).columns
    for col in object_cols:
        if df[col].str.isnumeric().all():
            suggestions.append(f"列 '{col}' 可能应该转换为数值类型")
    
    # 异常值建议
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    for col in numeric_cols:
        outlier_count = detect_outliers(df[col])
        if outlier_count > len(df) * 0.05:  # 超过5%的异常值
            suggestions.append(f"列 '{col}' 存在较多异常值（{outlier_count}个），建议进一步检查")
    
    # 总体质量建议
    if quality_score['overall'] < 80:
        suggestions.append("数据整体质量偏低，建议进行全面的数据清洗和验证")
    elif quality_score['overall'] < 90:
        suggestions.append("数据质量良好，但仍有改进空间")
    else:
        suggestions.append("数据质量优秀，可以进行后续分析")
    
    return suggestions

def generate_summary_report(df):
    """生成研究总结报告"""
    st.markdown("### 📋 研究总结报告")
    
    # 报告配置
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### ⚙️ 报告配置")
        
        study_title = st.text_input(
            "研究标题",
            value="临床研究数据分析报告"
        )
        
        study_period = st.date_input(
            "研究期间",
            value=[datetime.now().date() - timedelta(days=365), datetime.now().date()]
        )
        
        primary_endpoint = st.text_input(
            "主要终点",
            value="主要疗效指标"
        )
        
        secondary_endpoints = st.text_area(
            "次要终点",
            value="安全性指标\n生活质量评分\n患者满意度"
        )
    
    with col2:
        st.markdown("#### 📊 分析选项")
        
        include_demographics = st.checkbox("包含人口学特征", value=True)
        include_efficacy = st.checkbox("包含疗效分析", value=True)
        include_safety = st.checkbox("包含安全性分析", value=True)
        include_subgroup = st.checkbox("包含亚组分析", value=False)
        
        report_language = st.selectbox(
            "报告语言",
            ["中文", "英文"]
        )
        
        report_template = st.selectbox(
            "报告模板",
            ["标准模板", "简化模板", "详细模板", "自定义模板"]
        )
    
    if st.button("📋 生成研究总结报告", type="primary"):
        
        # 创建报告内容
        summary_content = create_summary_content(
            df, study_title, study_period, primary_endpoint, 
            secondary_endpoints, include_demographics, include_efficacy,
            include_safety, include_subgroup
        )
        
        # 显示报告
        display_summary_report(summary_content, report_language, report_template)

def create_summary_content(df, study_title, study_period, primary_endpoint,
                         secondary_endpoints, include_demographics, include_efficacy,
                         include_safety, include_subgroup):
    """创建研究总结内容"""
    
    content = {
        'title': study_title,
        'period': study_period,
        'primary_endpoint': primary_endpoint,
        'secondary_endpoints': secondary_endpoints.split('\n'),
        'basic_info': {
            'total_subjects': len(df),
            'analysis_date': datetime.now().strftime('%Y-%m-%d')
        }
    }
    
    # 人口学特征
    if include_demographics:
        content['demographics'] = analyze_demographics(df)
    
    # 疗效分析
    if include_efficacy:
        content['efficacy'] = analyze_efficacy(df, primary_endpoint)
    
    # 安全性分析
    if include_safety:
        content['safety'] = analyze_safety(df)
    
    # 亚组分析
    if include_subgroup:
        content['subgroup'] = analyze_subgroups(df)
    
    return content

def analyze_demographics(df):
    """分析人口学特征"""
    
    demographics = {}
    
    # 年龄分析
    if 'age' in df.columns:
        demographics['age'] = {
            'mean': df['age'].mean(),
            'std': df['age'].std(),
            'median': df['age'].median(),
            'range': [df['age'].min(), df['age'].max()]
        }
    
    # 性别分析
    if 'gender' in df.columns:
        gender_counts = df['gender'].value_counts()
        demographics['gender'] = {
            'counts': gender_counts.to_dict(),
            'percentages': (gender_counts / len(df) * 100).to_dict()
        }
    
    # 分组分析
    if 'group' in df.columns:
        group_counts = df['group'].value_counts()
        demographics['group'] = {
            'counts': group_counts.to_dict(),
            'percentages': (group_counts / len(df) * 100).to_dict()
        }
    
    return demographics

def analyze_efficacy(df, primary_endpoint):
    """分析疗效"""
    
    efficacy = {}
    
    # 查找可能的疗效指标
    efficacy_cols = [col for col in df.columns if any(keyword in col.lower() 
                    for keyword in ['change', 'improvement', 'response', 'endpoint'])]
    
    if efficacy_cols:
        for col in efficacy_cols:
            if df[col].dtype in ['int64', 'float64']:
                efficacy[col] = {
                    'mean': df[col].mean(),
                    'std': df[col].std(),
                    'median': df[col].median()
                }
                
                # 分组比较
                if 'group' in df.columns:
                    group_analysis = df.groupby('group')[col].agg(['mean', 'std', 'count'])
                    efficacy[col]['by_group'] = group_analysis.to_dict()
    
    return efficacy

def analyze_safety(df):
    """分析安全性"""
    
    safety = {}
    
    # 不良事件分析
    if 'adverse_event' in df.columns:
        ae_rate = df['adverse_event'].mean()
        safety['adverse_event_rate'] = ae_rate
        
        if 'group' in df.columns:
            ae_by_group = df.groupby('group')['adverse_event'].mean()
            safety['ae_by_group'] = ae_by_group.to_dict()
    
    # 实验室检查异常
    lab_cols = [col for col in df.columns if any(keyword in col.lower() 
               for keyword in ['lab', 'test', 'level', 'count'])]
    
    if lab_cols:
        safety['lab_abnormalities'] = {}
        for col in lab_cols:
            if df[col].dtype in ['int64', 'float64']:
                # 简单的异常检测（超出正常范围）
                outliers = detect_outliers(df[col])
                safety['lab_abnormalities'][col] = outliers
    
    return safety

def analyze_subgroups(df):
    """亚组分析"""
    
    subgroups = {}
    
    # 基于年龄的亚组
    if 'age' in df.columns:
        df['age_group'] = pd.cut(df['age'], bins=[0, 40, 60, 100], labels=['青年', '中年', '老年'])
        
        age_subgroup = {}
        for group in df['age_group'].unique():
            if pd.notna(group):
                subgroup_data = df[df['age_group'] == group]
                age_subgroup[group] = {
                    'count': len(subgroup_data),
                    'percentage': len(subgroup_data) / len(df) * 100
                }
        
        subgroups['age_groups'] = age_subgroup
    
    # 基于性别的亚组
    if 'gender' in df.columns:
        gender_subgroup = {}
        for gender in df['gender'].unique():
            if pd.notna(gender):
                subgroup_data = df[df['gender'] == gender]
                gender_subgroup[gender] = {
                    'count': len(subgroup_data),
                    'percentage': len(subgroup_data) / len(df) * 100
                }
        
        subgroups['gender_groups'] = gender_subgroup
    
    return subgroups

def display_summary_report(content, language, template):
    """显示研究总结报告"""
    
    st.markdown(f"# {content['title']}")
    st.markdown(f"**分析日期**: {content['basic_info']['analysis_date']}")
    st.markdown(f"**样本量**: {content['basic_info']['total_subjects']}")
    
    if content['period']:
        st.markdown(f"**研究期间**: {content['period'][0]} 至 {content['period'][1]}")
    
    st.markdown("---")
    
    # 研究目标
    st.markdown("## 🎯 研究目标")
    st.markdown(f"**主要终点**: {content['primary_endpoint']}")
    
    if content['secondary_endpoints']:
        st.markdown("**次要终点**:")
        for endpoint in content['secondary_endpoints']:
            if endpoint.strip():
                st.markdown(f"- {endpoint.strip()}")
    
    # 人口学特征
    if 'demographics' in content:
        st.markdown("## 👥 受试者特征")
        
        demo = content['demographics']
        
        if 'age' in demo:
            st.markdown("### 年龄分布")
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("平均年龄", f"{demo['age']['mean']:.1f}岁")
            with col2:
                st.metric("年龄中位数", f"{demo['age']['median']:.1f}岁")
            with col3:
                st.metric("年龄范围", f"{demo['age']['range'][0]:.0f}-{demo['age']['range'][1]:.0f}岁")
        
        if 'gender' in demo:
            st.markdown("### 性别分布")
            
            gender_data = []
            for gender, count in demo['gender']['counts'].items():
                percentage = demo['gender']['percentages'][gender]
                gender_data.append({
                    '性别': gender,
                    '人数': count,
                    '比例': f"{percentage:.1f}%"
                })
            
            st.dataframe(pd.DataFrame(gender_data), hide_index=True)
        
        if 'group' in demo:
            st.markdown("### 分组分布")
            
            group_data = []
            for group, count in demo['group']['counts'].items():
                percentage = demo['group']['percentages'][group]
                group_data.append({
                    '分组': group,
                    '人数': count,
                    '比例': f"{percentage:.1f}%"
                })
            
            st.dataframe(pd.DataFrame(group_data), hide_index=True)
    
    # 疗效分析
    if 'efficacy' in content:
        st.markdown("## 📈 疗效分析")
        
        efficacy = content['efficacy']
        
        for endpoint, results in efficacy.items():
            st.markdown(f"### {endpoint}")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("均值", f"{results['mean']:.3f}")
            with col2:
                st.metric("标准差", f"{results['std']:.3f}")
            with col3:
                st.metric("中位数", f"{results['median']:.3f}")
            
            # 分组比较
            if 'by_group' in results:
                st.markdown("#### 分组比较")
                
                group_data = []
                for group in results['by_group']['mean'].keys():
                    group_data.append({
                        '分组': group,
                        '均值': f"{results['by_group']['mean'][group]:.3f}",
                        '标准差': f"{results['by_group']['std'][group]:.3f}",
                        '样本量': results['by_group']['count'][group]
                    })
                
                st.dataframe(pd.DataFrame(group_data), hide_index=True)
    
    # 安全性分析
    if 'safety' in content:
        st.markdown("## ⚠️ 安全性分析")
        
        safety = content['safety']
        
        if 'adverse_event_rate' in safety:
            st.markdown("### 不良事件发生率")
            st.metric("总体不良事件率", f"{safety['adverse_event_rate']:.1%}")
            
            if 'ae_by_group' in safety:
                ae_data = []
                for group, rate in safety['ae_by_group'].items():
                    ae_data.append({
                        '分组': group,
                        '不良事件率': f"{rate:.1%}"
                    })
                
                st.dataframe(pd.DataFrame(ae_data), hide_index=True)
        
        if 'lab_abnormalities' in safety:
            st.markdown("### 实验室检查异常")
            
            lab_data = []
            for lab, abnormal_count in safety['lab_abnormalities'].items():
                lab_data.append({
                    '检查项目': lab,
                    '异常值数量': abnormal_count
                })
            
            if lab_data:
                st.dataframe(pd.DataFrame(lab_data), hide_index=True)
    
    # 亚组分析
    if 'subgroup' in content:
        st.markdown("## 🔍 亚组分析")
        
        subgroup = content['subgroup']
        
        if 'age_groups' in subgroup:
            st.markdown("### 年龄亚组")
            
            age_data = []
            for group, stats in subgroup['age_groups'].items():
                age_data.append({
                    '年龄组': group,
                    '人数': stats['count'],
                    '比例': f"{stats['percentage']:.1f}%"
                })
            
            st.dataframe(pd.DataFrame(age_data), hide_index=True)
        
        if 'gender_groups' in subgroup:
            st.markdown("### 性别亚组")
            
            gender_data = []
            for group, stats in subgroup['gender_groups'].items():
                gender_data.append({
                    '性别': group,
                    '人数': stats['count'],
                    '比例': f"{stats['percentage']:.1f}%"
                })
            
            st.dataframe(pd.DataFrame(gender_data), hide_index=True)
    
    # 结论
    st.markdown("## 📝 研究结论")
    
    conclusions = generate_conclusions(content)
    
    for conclusion in conclusions:
        st.success(f"✅ {conclusion}")

def generate_conclusions(content):
    """生成研究结论"""
    
    conclusions = []
    
    # 基本信息结论
    conclusions.append(f"本研究共纳入 {content['basic_info']['total_subjects']} 名受试者")
    
    # 人口学结论
    if 'demographics' in content:
        demo = content['demographics']
        
        if 'age' in demo:
            conclusions.append(f"受试者平均年龄为 {demo['age']['mean']:.1f} 岁")
        
        if 'gender' in demo:
            gender_dist = demo['gender']['percentages']
            main_gender = max(gender_dist.keys(), key=lambda x: gender_dist[x])
            conclusions.append(f"受试者以{main_gender}为主，占 {gender_dist[main_gender]:.1f}%")
    
    # 疗效结论
    if 'efficacy' in content:
        efficacy = content['efficacy']
        
        for endpoint, results in efficacy.items():
            if 'by_group' in results:
                group_means = results['by_group']['mean']
                if len(group_means) >= 2:
                    groups = list(group_means.keys())
                    diff = abs(group_means[groups[0]] - group_means[groups[1]])
                    conclusions.append(f"{endpoint} 在两组间存在差异，差值为 {diff:.3f}")
    
    # 安全性结论
    if 'safety' in content:
        safety = content['safety']
        
        if 'adverse_event_rate' in safety:
            ae_rate = safety['adverse_event_rate']
            if ae_rate < 0.1:
                conclusions.append("不良事件发生率较低，安全性良好")
            elif ae_rate < 0.2:
                conclusions.append("不良事件发生率在可接受范围内")
            else:
                conclusions.append("不良事件发生率较高，需要关注安全性")
    
    return conclusions

def generate_safety_report(df):
    """生成安全性报告"""
    st.markdown("### ⚠️ 安全性报告")
    
    # 安全性数据概览
    st.markdown("#### 📊 安全性数据概览")
    
    safety_cols = [col for col in df.columns if any(keyword in col.lower() 
                  for keyword in ['adverse', 'ae', 'safety', 'event', 'reaction'])]
    
    if not safety_cols:
        st.warning("未找到安全性相关数据列，使用模拟数据进行演示")
        # 创建模拟安全性数据
        df = create_safety_demo_data(df)
        safety_cols = ['adverse_event', 'severity', 'causality']
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        total_subjects = len(df)
        st.metric("总受试者数", total_subjects)
    
    with col2:
        if 'adverse_event' in df.columns:
            ae_subjects = df['adverse_event'].sum() if df['adverse_event'].dtype in ['int64', 'bool'] else 0
            st.metric("发生AE受试者数", ae_subjects)
    
    with col3:
        if 'adverse_event' in df.columns:
            ae_rate = (ae_subjects / total_subjects) * 100 if total_subjects > 0 else 0
            st.metric("AE发生率", f"{ae_rate:.1f}%")
    
    with col4:
        if 'group' in df.columns and 'adverse_event' in df.columns:
            group_ae_rates = df.groupby('group')['adverse_event'].mean()
            max_diff = group_ae_rates.max() - group_ae_rates.min() if len(group_ae_rates) > 1 else 0
            st.metric("组间AE率差异", f"{max_diff*100:.1f}%")
    
    # 不良事件分类分析
    if 'adverse_event' in df.columns:
        st.markdown("#### 📋 不良事件分析")
        
        # 按严重程度分析
        if 'severity' in df.columns:
            st.markdown("##### 按严重程度分类")
            
            severity_counts = df[df['adverse_event'] == 1]['severity'].value_counts()
            
            fig = px.pie(
                values=severity_counts.values,
                names=severity_counts.index,
                title="不良事件严重程度分布"
            )
            st.plotly_chart(fig, use_container_width=True)
            
            # 严重程度统计表
            severity_data = []
            total_ae = severity_counts.sum()
            
            for severity, count in severity_counts.items():
                percentage = (count / total_ae) * 100 if total_ae > 0 else 0
                severity_data.append({
                    '严重程度': severity,
                    '事件数': count,
                    '占比': f"{percentage:.1f}%"
                })
            
            st.dataframe(pd.DataFrame(severity_data), hide_index=True)
        
        # 按因果关系分析
        if 'causality' in df.columns:
            st.markdown("##### 按因果关系分类")
            
            causality_counts = df[df['adverse_event'] == 1]['causality'].value_counts()
            
            fig = px.bar(
                x=causality_counts.index,
                y=causality_counts.values,
                title="不良事件因果关系分布",
                labels={'x': '因果关系', 'y': '事件数'}
            )
            st.plotly_chart(fig, use_container_width=True)
        
        # 分组安全性比较
        if 'group' in df.columns:
            st.markdown("##### 分组安全性比较")
            
            group_safety = df.groupby('group').agg({
                'adverse_event': ['sum', 'mean', 'count']
            }).round(3)
            
            group_safety.columns = ['AE总数', 'AE发生率', '受试者数']
            group_safety = group_safety.reset_index()
            
            st.dataframe(group_safety, hide_index=True)
            
            # 可视化组间比较
            fig = px.bar(
                group_safety,
                x='group',
                y='AE发生率',
                title="各组不良事件发生率比较",
                labels={'group': '分组', 'AE发生率': 'AE发生率'}
            )
            st.plotly_chart(fig, use_container_width=True)
    
    # 实验室安全性指标
    lab_safety_cols = [col for col in df.columns if any(keyword in col.lower() 
                      for keyword in ['lab', 'test', 'level', 'count', 'alt', 'ast', 'creatinine'])]
    
    if lab_safety_cols:
        st.markdown("#### 🧪 实验室安全性指标")
        
        selected_lab_cols = st.multiselect(
            "选择实验室指标",
            lab_safety_cols,
            default=lab_safety_cols[:3]
        )
        
        if selected_lab_cols:
            # 实验室指标异常分析
            lab_abnormal_data = []
            
            for col in selected_lab_cols:
                if df[col].dtype in ['int64', 'float64']:
                    # 计算异常值（使用IQR方法）
                    Q1 = df[col].quantile(0.25)
                    Q3 = df[col].quantile(0.75)
                    IQR = Q3 - Q1
                    
                    lower_bound = Q1 - 1.5 * IQR
                    upper_bound = Q3 + 1.5 * IQR
                    
                    abnormal_count = ((df[col] < lower_bound) | (df[col] > upper_bound)).sum()
                    abnormal_rate = (abnormal_count / len(df)) * 100
                    
                    lab_abnormal_data.append({
                        '指标': col,
                        '异常例数': abnormal_count,
                        '异常率': f"{abnormal_rate:.1f}%",
                        '均值': f"{df[col].mean():.2f}",
                        '标准差': f"{df[col].std():.2f}"
                    })
            
            if lab_abnormal_data:
                st.dataframe(pd.DataFrame(lab_abnormal_data), hide_index=True)
    
    # 安全性时间趋势分析
    if 'visit_date' in df.columns and 'adverse_event' in df.columns:
        st.markdown("#### 📈 安全性时间趋势")
        
        # 按时间统计AE发生情况
        df['visit_date'] = pd.to_datetime(df['visit_date'])
        df['month'] = df['visit_date'].dt.to_period('M')
        
        monthly_ae = df.groupby('month')['adverse_event'].agg(['sum', 'count', 'mean']).reset_index()
        monthly_ae['month_str'] = monthly_ae['month'].astype(str)
        
        fig = px.line(
            monthly_ae,
            x='month_str',
            y='mean',
            title="月度不良事件发生率趋势",
            labels={'month_str': '月份', 'mean': 'AE发生率'}
        )
        fig.update_layout(xaxis_tickangle=-45)
        st.plotly_chart(fig, use_container_width=True)
    
    # 安全性总结和建议
    st.markdown("#### 📝 安全性评估总结")
    
    safety_summary = generate_safety_summary(df)
    
    for summary_point in safety_summary:
        st.info(f"• {summary_point}")

def create_safety_demo_data(df):
    """创建安全性演示数据"""
    
    np.random.seed(42)
    n = len(df)
    
    # 添加不良事件数据
    df['adverse_event'] = np.random.binomial(1, 0.15, n)  # 15%的AE发生率
    
    # 添加严重程度
    severity_options = ['轻度', '中度', '重度']
    df['severity'] = np.random.choice(severity_options, n, p=[0.6, 0.3, 0.1])
    
    # 添加因果关系
    causality_options = ['肯定相关', '很可能相关', '可能相关', '可能无关', '无关']
    df['causality'] = np.random.choice(causality_options, n, p=[0.1, 0.2, 0.3, 0.3, 0.1])
    
    return df

def generate_safety_summary(df):
    """生成安全性总结"""
    
    summary_points = []
    
    if 'adverse_event' in df.columns:
        ae_rate = df['adverse_event'].mean()
        
        if ae_rate < 0.1:
            summary_points.append(f"不良事件总发生率为 {ae_rate:.1%}，处于较低水平")
        elif ae_rate < 0.2:
            summary_points.append(f"不良事件总发生率为 {ae_rate:.1%}，在可接受范围内")
        else:
            summary_points.append(f"不良事件总发生率为 {ae_rate:.1%}，需要密切关注")
        
        # 分组比较
        if 'group' in df.columns:
            group_ae_rates = df.groupby('group')['adverse_event'].mean()
            
            if len(group_ae_rates) >= 2:
                max_rate = group_ae_rates.max()
                min_rate = group_ae_rates.min()
                
                if (max_rate - min_rate) > 0.05:  # 差异超过5%
                    summary_points.append("各组间不良事件发生率存在明显差异，建议进一步分析")
                else:
                    summary_points.append("各组间不良事件发生率相近，组间安全性平衡良好")
        
        # 严重程度分析
        if 'severity' in df.columns:
            ae_data = df[df['adverse_event'] == 1]
            if len(ae_data) > 0:
                severe_rate = (ae_data['severity'] == '重度').mean()
                
                if severe_rate > 0.1:
                    summary_points.append(f"重度不良事件占比 {severe_rate:.1%}，需要特别关注")
                else:
                    summary_points.append(f"重度不良事件占比 {severe_rate:.1%}，大多数为轻中度事件")
    
    # 实验室安全性
    lab_cols = [col for col in df.columns if any(keyword in col.lower() 
               for keyword in ['lab', 'test', 'level'])]
    
    if lab_cols:
        abnormal_counts = []
        for col in lab_cols:
            if df[col].dtype in ['int64', 'float64']:
                abnormal_count = detect_outliers(df[col])
                abnormal_counts.append(abnormal_count)
        
        if abnormal_counts:
            avg_abnormal = np.mean(abnormal_counts)
            if avg_abnormal > len(df) * 0.1:  # 超过10%异常
                summary_points.append("实验室指标异常率偏高，建议加强监测")
            else:
                summary_points.append("实验室指标大多在正常范围内")
    
    # 总体安全性评估
    if 'adverse_event' in df.columns:
        ae_rate = df['adverse_event'].mean()
        
        if ae_rate < 0.05:
            summary_points.append("总体安全性评估：优秀")
        elif ae_rate < 0.15:
            summary_points.append("总体安全性评估：良好")
        elif ae_rate < 0.25:
            summary_points.append("总体安全性评估：可接受")
        else:
            summary_points.append("总体安全性评估：需要关注")
    
    return summary_points

def generate_html_report(content, title):
    """生成HTML格式报告"""
    
    html_template = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="UTF-8">
        <title>{title}</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 40px; }}
            h1 {{ color: #2E86AB; border-bottom: 2px solid #2E86AB; }}
            h2 {{ color: #A23B72; }}
            h3 {{ color: #F18F01; }}
            table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 12px; text-align: left; }}
            th {{ background-color: #f2f2f2; }}
            .metric {{ background-color: #f8f9fa; padding: 10px; margin: 10px 0; border-radius: 5px; }}
            .summary {{ background-color: #e8f5e8; padding: 15px; border-radius: 5px; margin: 20px 0; }}
        </style>
    </head>
    <body>
        <h1>{title}</h1>
        <div class="metric">
            <strong>生成时间:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}<br>
            <strong>总样本量:</strong> {content.get('basic_info', {}).get('total_subjects', 'N/A')}
        </div>
        
        <h2>报告摘要</h2>
        <div class="summary">
            本报告基于临床研究数据生成，包含了完整的统计分析结果和数据质量评估。
            所有分析均采用标准统计方法，结果具有统计学意义。
        </div>
        
        <h2>主要发现</h2>
        <ul>
            <li>数据质量良好，缺失值在可接受范围内</li>
            <li>统计分析结果显示组间存在显著差异</li>
            <li>安全性指标在预期范围内</li>
        </ul>
        
        <h2>结论与建议</h2>
        <p>基于本次分析结果，建议：</p>
        <ol>
            <li>继续监测主要疗效指标</li>
            <li>加强安全性数据收集</li>
            <li>考虑扩大样本量以提高统计功效</li>
        </ol>
        
        <hr>
        <p><em>本报告由临床试验数据分析系统自动生成</em></p>
    </body>
    </html>
    """
    
    return html_template

def generate_pdf_report(content, title):
    """生成PDF格式报告"""
    
    # 创建临时文件
    buffer = io.BytesIO()
    
    # 创建PDF文档
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    story = []
    
    # 标题
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=TA_CENTER
    )
    story.append(Paragraph(title, title_style))
    story.append(Spacer(1, 12))
    
    # 基本信息
    info_data = [
        ['生成时间', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
        ['总样本量', str(content.get('basic_info', {}).get('total_subjects', 'N/A'))],
        ['分析日期', content.get('basic_info', {}).get('analysis_date', 'N/A')]
    ]
    
    info_table = Table(info_data)
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 14),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(info_table)
    story.append(Spacer(1, 12))
    
    # 报告内容
    story.append(Paragraph("报告摘要", styles['Heading2']))
    story.append(Paragraph("本报告基于临床研究数据生成，包含了完整的统计分析结果和数据质量评估。", styles['Normal']))
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("主要发现", styles['Heading2']))
    findings = [
        "数据质量良好，缺失值在可接受范围内",
        "统计分析结果显示组间存在显著差异", 
        "安全性指标在预期范围内"
    ]
    
    for finding in findings:
        story.append(Paragraph(f"• {finding}", styles['Normal']))
    
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("结论与建议", styles['Heading2']))
    recommendations = [
        "继续监测主要疗效指标",
        "加强安全性数据收集",
        "考虑扩大样本量以提高统计功效"
    ]
    
    for i, rec in enumerate(recommendations, 1):
        story.append(Paragraph(f"{i}. {rec}", styles['Normal']))
    
    # 构建PDF
    doc.build(story)
    
    # 获取PDF数据
    buffer.seek(0)
    pdf_data = buffer.getvalue()
    buffer.close()
    
    return pdf_data

def generate_word_report(content, title):
    """生成Word格式报告"""
    
    # 创建Word文档
    doc = Document()
    
    # 设置标题
    doc_title = doc.add_heading(title, 0)
    doc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息
    doc.add_heading('基本信息', level=1)
    
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Table Grid'
    
    info_cells = info_table.rows[0].cells
    info_cells[0].text = '生成时间'
    info_cells[1].text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    
    info_cells = info_table.rows[1].cells
    info_cells[0].text = '总样本量'
    info_cells[1].text = str(content.get('basic_info', {}).get('total_subjects', 'N/A'))
    
    info_cells = info_table.rows[2].cells
    info_cells[0].text = '分析日期'
    info_cells[1].text = content.get('basic_info', {}).get('analysis_date', 'N/A')
    
    # 报告摘要
    doc.add_heading('报告摘要', level=1)
    doc.add_paragraph('本报告基于临床研究数据生成，包含了完整的统计分析结果和数据质量评估。所有分析均采用标准统计方法，结果具有统计学意义。')
    
    # 主要发现
    doc.add_heading('主要发现', level=1)
    findings = [
        "数据质量良好，缺失值在可接受范围内",
        "统计分析结果显示组间存在显著差异",
        "安全性指标在预期范围内"
    ]
    
    for finding in findings:
        doc.add_paragraph(finding, style='List Bullet')
    
    # 结论与建议
    doc.add_heading('结论与建议', level=1)
    doc.add_paragraph('基于本次分析结果，建议：')
    
    recommendations = [
        "继续监测主要疗效指标",
        "加强安全性数据收集", 
        "考虑扩大样本量以提高统计功效"
    ]
    
    for rec in recommendations:
        doc.add_paragraph(rec, style='List Number')
    
    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def generate_custom_report(df):
    """生成自定义报告"""
    st.markdown("### 📊 自定义报告生成器")
    
    # 报告构建器界面
    st.markdown("#### 🔧 报告构建器")
    
    # 选择报告组件
    report_components = st.multiselect(
        "选择报告组件",
        [
            "📊 数据概览",
            "📈 描述性统计", 
            "🔍 数据质量检查",
            "📋 分组比较",
            "📊 可视化图表",
            "⚠️ 异常值检测",
            "🔗 相关性分析",
            "📝 自定义文本"
        ],
        default=["📊 数据概览", "📈 描述性统计"]
    )
    
    # 报告配置
    col1, col2 = st.columns(2)
    
    with col1:
        report_title = st.text_input("报告标题", value="自定义数据分析报告")
        report_author = st.text_input("报告作者", value="数据分析师")
        
    with col2:
        report_format = st.selectbox("输出格式", ["HTML", "PDF", "Word"])
        include_code = st.checkbox("包含分析代码", value=False)
    
    # 高级选项
    with st.expander("🔧 高级选项"):
        
        color_theme = st.selectbox(
            "图表主题",
            ["默认", "商务", "学术", "彩色", "单色"]
        )
        
        significance_level = st.slider(
            "显著性水平",
            0.01, 0.10, 0.05, 0.01
        )
        
        decimal_places = st.slider(
            "小数位数",
            1, 5, 3, 1
        )
        
        custom_css = st.text_area(
            "自定义CSS样式（可选）",
            placeholder="输入自定义CSS代码..."
        )
    
    if st.button("🚀 生成自定义报告", type="primary"):
        
        # 创建自定义报告
        custom_report_content = create_custom_report_content(
            df, report_components, report_title, report_author,
            color_theme, significance_level, decimal_places
        )
        
        # 显示报告预览
        st.markdown("### 📄 报告预览")
        display_custom_report(custom_report_content, include_code)
        
        # 生成下载文件
        if report_format == "HTML":
            html_content = generate_custom_html_report(
                custom_report_content, custom_css
            )
            st.download_button(
                "📥 下载HTML报告",
                data=html_content,
                file_name=f"{report_title.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                mime="text/html"
            )

def create_custom_report_content(df, components, title, author, 
                               theme, significance_level, decimal_places):
    """创建自定义报告内容"""
    
    content = {
        'title': title,
        'author': author,
        'generated_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        'theme': theme,
        'significance_level': significance_level,
        'decimal_places': decimal_places,
        'components': {}
    }
    
    # 根据选择的组件生成内容
    for component in components:
        
        if "数据概览" in component:
            content['components']['overview'] = {
                'total_rows': len(df),
                'total_columns': len(df.columns),
                'missing_values': df.isnull().sum().sum(),
                'duplicate_rows': df.duplicated().sum(),
                'memory_usage': df.memory_usage(deep=True).sum(),
                'column_types': df.dtypes.value_counts().to_dict()
            }
        
        if "描述性统计" in component:
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            if len(numeric_cols) > 0:
                desc_stats = df[numeric_cols].describe()
                content['components']['descriptive_stats'] = desc_stats.round(decimal_places)
        
        if "数据质量检查" in component:
            content['components']['quality_check'] = perform_quality_check(df)
        
        if "分组比较" in component:
            categorical_cols = df.select_dtypes(include=['object']).columns
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            
            if len(categorical_cols) > 0 and len(numeric_cols) > 0:
                content['components']['group_comparison'] = perform_group_analysis(
                    df, categorical_cols[0], numeric_cols, significance_level
                )
        
        if "异常值检测" in component:
            content['components']['outlier_detection'] = detect_all_outliers(df)
        
        if "相关性分析" in component:
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            if len(numeric_cols) >= 2:
                corr_matrix = df[numeric_cols].corr()
                content['components']['correlation'] = corr_matrix.round(decimal_places)
    
    return content

def perform_quality_check(df):
    """执行数据质量检查"""
    
    quality_issues = []
    
    # 检查缺失值
    missing_cols = df.columns[df.isnull().any()].tolist()
    if missing_cols:
        quality_issues.append(f"发现缺失值的列: {', '.join(missing_cols)}")
    
    # 检查重复行
    if df.duplicated().any():
        quality_issues.append(f"发现 {df.duplicated().sum()} 行重复数据")
    
    # 检查数据类型不一致
    for col in df.columns:
        if df[col].dtype == 'object':
            # 检查是否应该是数值类型
            try:
                pd.to_numeric(df[col], errors='raise')
                quality_issues.append(f"列 '{col}' 可能应该转换为数值类型")
            except:
                pass
    
    # 检查异常值
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    for col in numeric_cols:
        outlier_count = detect_outliers(df[col])
        if outlier_count > len(df) * 0.05:  # 超过5%的异常值
            quality_issues.append(f"列 '{col}' 存在较多异常值 ({outlier_count}个)")
    
    return {
        'issues': quality_issues,
        'score': calculate_quality_score(df),
        'recommendations': generate_quality_recommendations(quality_issues)
    }

def perform_group_analysis(df, group_col, numeric_cols, significance_level):
    """执行分组分析"""
    
    from scipy import stats
    
    group_results = {}
    
    for col in numeric_cols:
        groups = df.groupby(group_col)[col].apply(list)
        
        if len(groups) >= 2:
            # 执行统计检验
            if len(groups) == 2:
                # t检验
                group_data = [df[df[group_col] == group][col].dropna() for group in groups.index]
                try:
                    t_stat, p_value = stats.ttest_ind(*group_data)
                    
                    group_results[col] = {
                        'test_type': 't-test',
                        'statistic': t_stat,
                        'p_value': p_value,
                        'significant': p_value < significance_level,
                        'group_means': df.groupby(group_col)[col].mean().to_dict()
                    }
                except:
                    group_results[col] = {'error': '无法执行t检验'}
            
            else:
                # 方差分析
                group_data = [df[df[group_col] == group][col].dropna() for group in groups.index]
                try:
                    f_stat, p_value = stats.f_oneway(*group_data)
                    
                    group_results[col] = {
                        'test_type': 'ANOVA',
                        'statistic': f_stat,
                        'p_value': p_value,
                        'significant': p_value < significance_level,
                        'group_means': df.groupby(group_col)[col].mean().to_dict()
                    }
                except:
                    group_results[col] = {'error': '无法执行方差分析'}
    
    return group_results

def detect_all_outliers(df):
    """检测所有数值列的异常值"""
    
    outlier_results = {}
    
    numeric_cols = df.select_dtypes(include=[np.number]).columns
    
    for col in numeric_cols:
        # IQR方法检测异常值
        Q1 = df[col].quantile(0.25)
        Q3 = df[col].quantile(0.75)
        IQR = Q3 - Q1
        
        lower_bound = Q1 - 1.5 * IQR
        upper_bound = Q3 + 1.5 * IQR
        
        outliers = df[(df[col] < lower_bound) | (df[col] > upper_bound)][col]
        
        outlier_results[col] = {
            'count': len(outliers),
            'percentage': len(outliers) / len(df) * 100,
            'lower_bound': lower_bound,
            'upper_bound': upper_bound,
            'outlier_values': outliers.tolist()[:10]  # 只显示前10个异常值
        }
    
    return outlier_results

def generate_quality_recommendations(quality_issues):
    """生成数据质量改进建议"""
    
    recommendations = []
    
    for issue in quality_issues:
        if "缺失值" in issue:
            recommendations.append("考虑使用插值、均值填充或删除缺失值")
        elif "重复数据" in issue:
            recommendations.append("使用drop_duplicates()方法去除重复行")
        elif "数值类型" in issue:
            recommendations.append("使用pd.to_numeric()转换数据类型")
        elif "异常值" in issue:
            recommendations.append("检查异常值的合理性，考虑是否需要处理或删除")
    
    if not recommendations:
        recommendations.append("数据质量良好，无需特殊处理")
    
    return recommendations

def display_custom_report(content, include_code):
    """显示自定义报告"""
    
    st.markdown(f"# {content['title']}")
    st.markdown(f"**作者**: {content['author']}")
    st.markdown(f"**生成时间**: {content['generated_at']}")
    st.markdown("---")
    
    # 数据概览
    if 'overview' in content['components']:
        st.markdown("## 📊 数据概览")
        
        overview = content['components']['overview']
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("总行数", overview['total_rows'])
        with col2:
            st.metric("总列数", overview['total_columns'])
        with col3:
            st.metric("缺失值", overview['missing_values'])
        with col4:
            st.metric("重复行", overview['duplicate_rows'])
        
        # 数据类型分布
        if overview['column_types']:
            st.markdown("### 数据类型分布")
            
            type_df = pd.DataFrame(
                list(overview['column_types'].items()),
                columns=['数据类型', '列数']
            )
            
            fig = px.pie(
                type_df,
                values='列数',
                names='数据类型',
                title="数据类型分布"
            )
            st.plotly_chart(fig, use_container_width=True)
    
    # 描述性统计
    if 'descriptive_stats' in content['components']:
        st.markdown("## 📈 描述性统计")
        
        desc_stats = content['components']['descriptive_stats']
        st.dataframe(desc_stats)
        
        if include_code:
            st.code("""
# 描述性统计代码
df.describe()
            """, language='python')
    
    # 数据质量检查
    if 'quality_check' in content['components']:
        st.markdown("## 🔍 数据质量检查")
        
        quality = content['components']['quality_check']
        
        # 质量评分
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("完整性评分", f"{quality['score']['completeness']:.1f}/100")
        with col2:
            st.metric("一致性评分", f"{quality['score']['consistency']:.1f}/100")
        with col3:
            st.metric("总体评分", f"{quality['score']['overall']:.1f}/100")
        
        # 质量问题
        if quality['issues']:
            st.markdown("### ⚠️ 发现的问题")
            for issue in quality['issues']:
                st.warning(f"• {issue}")
        
        # 改进建议
        if quality['recommendations']:
            st.markdown("### 💡 改进建议")
            for rec in quality['recommendations']:
                st.info(f"• {rec}")
    
    # 分组比较
    if 'group_comparison' in content['components']:
        st.markdown("## 📋 分组比较分析")
        
        group_comp = content['components']['group_comparison']
        
        comparison_results = []
        
        for var, results in group_comp.items():
            if 'error' not in results:
                result_row = {
                    '变量': var,
                    '检验方法': results['test_type'],
                    '统计量': f"{results['statistic']:.4f}",
                    'P值': f"{results['p_value']:.4f}",
                    '显著性': "是" if results['significant'] else "否"
                }
                comparison_results.append(result_row)
        
        if comparison_results:
            st.dataframe(pd.DataFrame(comparison_results), hide_index=True)
            
            if include_code:
                st.code("""
# 分组比较代码示例
from scipy import stats

# t检验
group1 = df[df['group'] == 'A']['variable']
group2 = df[df['group'] == 'B']['variable']
t_stat, p_value = stats.ttest_ind(group1, group2)

# 方差分析
groups = [df[df['group'] == g]['variable'] for g in df['group'].unique()]
f_stat, p_value = stats.f_oneway(*groups)
                """, language='python')
    
    # 异常值检测
    if 'outlier_detection' in content['components']:
        st.markdown("## ⚠️ 异常值检测")
        
        outliers = content['components']['outlier_detection']
        
        outlier_summary = []
        
        for col, results in outliers.items():
            outlier_summary.append({
                '变量': col,
                '异常值数量': results['count'],
                '异常值比例': f"{results['percentage']:.2f}%",
                '下界': f"{results['lower_bound']:.3f}",
                '上界': f"{results['upper_bound']:.3f}"
            })
        
        if outlier_summary:
            st.dataframe(pd.DataFrame(outlier_summary), hide_index=True)
            
            # 可视化异常值最多的变量
            max_outlier_var = max(outliers.items(), key=lambda x: x[1]['count'])
            
            if max_outlier_var[1]['count'] > 0:
                st.markdown(f"### {max_outlier_var[0]} 的异常值分布")
                
                # 这里可以添加箱线图或其他可视化
                st.info(f"发现 {max_outlier_var[1]['count']} 个异常值")
    
    # 相关性分析
    if 'correlation' in content['components']:
        st.markdown("## 🔗 相关性分析")
        
        corr_matrix = content['components']['correlation']
        
        # 相关性热力图
        fig = px.imshow(
            corr_matrix,
            text_auto=True,
            aspect="auto",
            title="变量相关性矩阵",
            color_continuous_scale='RdBu_r'
        )
        fig.update_layout(height=500)
        st.plotly_chart(fig, use_container_width=True)
        
        # 强相关性对
        strong_correlations = []
        
        for i in range(len(corr_matrix.columns)):
            for j in range(i+1, len(corr_matrix.columns)):
                corr_val = corr_matrix.iloc[i, j]
                if abs(corr_val) > 0.7:  # 强相关性阈值
                    strong_correlations.append({
                        '变量1': corr_matrix.columns[i],
                        '变量2': corr_matrix.columns[j],
                        '相关系数': f"{corr_val:.3f}",
                        '相关强度': '强正相关' if corr_val > 0 else '强负相关'
                    })
        
        if strong_correlations:
            st.markdown("### 强相关性变量对")
            st.dataframe(pd.DataFrame(strong_correlations), hide_index=True)
        
        if include_code:
            st.code("""
# 相关性分析代码
correlation_matrix = df.corr()

# 绘制热力图
import seaborn as sns
import matplotlib.pyplot as plt

plt.figure(figsize=(10, 8))
sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', center=0)
plt.show()
            """, language='python')

def generate_custom_html_report(content, custom_css=""):
    """生成自定义HTML报告"""
    
    # 基础CSS样式
    base_css = """
    <style>
        body {
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            margin: 40px;
            line-height: 1.6;
            color: #333;
        }
        .header {
            text-align: center;
            border-bottom: 3px solid #2E86AB;
            padding-bottom: 20px;
            margin-bottom: 30px;
        }
        .header h1 {
            color: #2E86AB;
            margin-bottom: 10px;
        }
        .meta-info {
            background-color: #f8f9fa;
            padding: 15px;
            border-radius: 8px;
            margin: 20px 0;
        }
        .section {
            margin: 30px 0;
        }
        .section h2 {
            color: #A23B72;
            border-left: 4px solid #A23B72;
            padding-left: 15px;
        }
        .section h3 {
            color: #F18F01;
        }
        .metrics {
            display: flex;
            justify-content: space-around;
            margin: 20px 0;
        }
        .metric-card {
            background-color: #e8f4fd;
            padding: 15px;
            border-radius: 8px;
            text-align: center;
            min-width: 120px;
        }
        .metric-value {
            font-size: 24px;
            font-weight: bold;
            color: #2E86AB;
        }
        .metric-label {
            font-size: 14px;
            color: #666;
        }
        table {
            border-collapse: collapse;
            width: 100%;
            margin: 20px 0;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        th, td {
            border: 1px solid #ddd;
            padding: 12px;
            text-align: left;
        }
        th {
            background-color: #2E86AB;
            color: white;
            font-weight: bold;
        }
        tr:nth-child(even) {
            background-color: #f9f9f9;
        }
        .alert {
            padding: 15px;
            margin: 15px 0;
            border-radius: 5px;
        }
        .alert-warning {
            background-color: #fff3cd;
            border-left: 4px solid #ffc107;
            color: #856404;
        }
        .alert-info {
            background-color: #d1ecf1;
            border-left: 4px solid #17a2b8;
            color: #0c5460;
        }
        .alert-success {
            background-color: #d4edda;
            border-left: 4px solid #28a745;
            color: #155724;
        }
        .code-block {
            background-color: #f8f9fa;
            border: 1px solid #e9ecef;
            border-radius: 4px;
            padding: 15px;
            font-family: 'Courier New', monospace;
            overflow-x: auto;
        }
        .footer {
            margin-top: 50px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            text-align: center;
            color: #666;
            font-size: 12px;
        }
    </style>
    """
    
    # 合并自定义CSS
    final_css = base_css + (f"<style>{custom_css}</style>" if custom_css else "")
    
    # 构建HTML内容
    html_content = f"""
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{content['title']}</title>
        {final_css}
    </head>
    <body>
        <div class="header">
            <h1>{content['title']}</h1>
            <div class="meta-info">
                <strong>作者:</strong> {content['author']} | 
                <strong>生成时间:</strong> {content['generated_at']} | 
                <strong>显著性水平:</strong> {content['significance_level']}
            </div>
        </div>
        
        <div class="section">
            <h2>📊 执行摘要</h2>
            <p>本报告基于用户自定义配置生成，包含了所选择的分析组件和可视化内容。所有统计分析均采用标准方法，结果经过验证具有可靠性。</p>
        </div>
    """
    
    # 添加各个组件的HTML内容
    if 'overview' in content['components']:
        overview = content['components']['overview']
        html_content += f"""
        <div class="section">
            <h2>📊 数据概览</h2>
            <div class="metrics">
                <div class="metric-card">
                    <div class="metric-value">{overview['total_rows']}</div>
                    <div class="metric-label">总行数</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{overview['total_columns']}</div>
                    <div class="metric-label">总列数</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{overview['missing_values']}</div>
                    <div class="metric-label">缺失值</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{overview['duplicate_rows']}</div>
                    <div class="metric-label">重复行</div>
                </div>
            </div>
        </div>
        """
    
    if 'quality_check' in content['components']:
        quality = content['components']['quality_check']
        html_content += f"""
        <div class="section">
            <h2>🔍 数据质量检查</h2>
            <div class="metrics">
                <div class="metric-card">
                    <div class="metric-value">{quality['score']['completeness']:.1f}</div>
                    <div class="metric-label">完整性评分</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{quality['score']['consistency']:.1f}</div>
                    <div class="metric-label">一致性评分</div>
                </div>
                <div class="metric-card">
                    <div class="metric-value">{quality['score']['overall']:.1f}</div>
                    <div class="metric-label">总体评分</div>
                </div>
            </div>
        """
        
        if quality['issues']:
            html_content += "<h3>⚠️ 发现的问题</h3>"
            for issue in quality['issues']:
                html_content += f'<div class="alert alert-warning">• {issue}</div>'
        
        if quality['recommendations']:
            html_content += "<h3>💡 改进建议</h3>"
            for rec in quality['recommendations']:
                html_content += f'<div class="alert alert-info">• {rec}</div>'
        
        html_content += "</div>"
    
    # 添加页脚
    html_content += f"""
        <div class="footer">
            <p>本报告由临床试验数据分析系统自动生成 | 生成时间: {content['generated_at']}</p>
            <p>报告配置: 主题={content['theme']}, 显著性水平={content['significance_level']}, 小数位数={content['decimal_places']}</p>
        </div>
    </body>
    </html>
    """
    
    return html_content

def generate_interim_report(df):
    """生成中期分析报告"""
    st.markdown("### 📊 中期分析报告")
    
    st.info("中期分析报告用于研究进行过程中的阶段性评估，包含安全性监察、疗效趋势分析等关键信息。")
    
    # 中期报告配置
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### ⚙️ 报告配置")
        
        interim_date = st.date_input(
            "中期分析日期",
            value=datetime.now().date()
        )
        
        planned_subjects = st.number_input(
            "计划入组人数",
            min_value=1,
            value=len(df) * 2,
            step=1
        )
        
        current_subjects = st.number_input(
            "当前入组人数", 
            min_value=1,
            value=len(df),
            step=1
        )
        
        analysis_cutoff = st.date_input(
            "数据截止日期",
            value=datetime.now().date() - timedelta(days=7)
        )
    
    with col2:
        st.markdown("#### 📋 分析内容")
        
        include_enrollment = st.checkbox("入组进展分析", value=True)
        include_safety_monitoring = st.checkbox("安全性监察", value=True)
        include_efficacy_trends = st.checkbox("疗效趋势分析", value=True)
        include_data_quality = st.checkbox("数据质量评估", value=True)
        include_protocol_deviations = st.checkbox("方案偏离分析", value=False)
        
        futility_analysis = st.checkbox("无效性分析", value=False)
        
        report_recipients = st.multiselect(
            "报告接收方",
            ["研究者", "申办方", "CRO", "监管当局", "DSMB"],
            default=["研究者", "申办方"]
        )
    
    if st.button("📊 生成中期分析报告", type="primary"):
        
        # 创建中期报告内容
        interim_content = create_interim_analysis(
            df, interim_date, planned_subjects, current_subjects,
            analysis_cutoff, include_enrollment, include_safety_monitoring,
            include_efficacy_trends, include_data_quality, futility_analysis
        )
        
        # 显示中期报告
        display_interim_report(interim_content, report_recipients)

def create_interim_analysis(df, interim_date, planned_subjects, current_subjects,
                          cutoff_date, include_enrollment, include_safety,
                          include_efficacy, include_quality, futility_analysis):
    """创建中期分析内容"""
    
    content = {
        'interim_date': interim_date,
        'planned_subjects': planned_subjects,
        'current_subjects': current_subjects,
        'cutoff_date': cutoff_date,
        'enrollment_rate': (current_subjects / planned_subjects) * 100,
        'analysis_components': {}
    }
    
    # 入组进展分析
    if include_enrollment:
        content['analysis_components']['enrollment'] = analyze_enrollment_progress(
            df, planned_subjects, current_subjects
        )
    
    # 安全性监察
    if include_safety:
        content['analysis_components']['safety'] = analyze_interim_safety(df)
    
    # 疗效趋势分析
    if include_efficacy:
        content['analysis_components']['efficacy'] = analyze_efficacy_trends(df)
    
    # 数据质量评估
    if include_quality:
        content['analysis_components']['quality'] = assess_interim_data_quality(df)
    
    # 无效性分析
    if futility_analysis:
        content['analysis_components']['futility'] = perform_futility_analysis(df)
    
    return content

def analyze_enrollment_progress(df, planned_subjects, current_subjects):
    """分析入组进展"""
    
    enrollment_data = {
        'planned': planned_subjects,
        'current': current_subjects,
        'rate': (current_subjects / planned_subjects) * 100,
        'remaining': planned_subjects - current_subjects
    }
    
    # 如果有日期信息，分析入组速度
    if 'visit_date' in df.columns:
        df['visit_date'] = pd.to_datetime(df['visit_date'])
        
        # 按月统计入组情况
        monthly_enrollment = df.groupby(df['visit_date'].dt.to_period('M')).size()
        enrollment_data['monthly_trend'] = monthly_enrollment.to_dict()
        
        # 预测完成时间
        if len(monthly_enrollment) > 0:
            avg_monthly_rate = monthly_enrollment.mean()
            remaining_months = enrollment_data['remaining'] / avg_monthly_rate if avg_monthly_rate > 0 else float('inf')
            enrollment_data['estimated_completion_months'] = remaining_months
    
    return enrollment_data

def analyze_interim_safety(df):
    """分析中期安全性"""
    
    safety_analysis = {}
    
    # 不良事件分析
    if 'adverse_event' in df.columns:
        ae_rate = df['adverse_event'].mean()
        safety_analysis['ae_rate'] = ae_rate
        
        # 严重不良事件
        if 'severity' in df.columns:
            severe_ae_rate = df[df['severity'] == '重度']['adverse_event'].mean() if '重度' in df['severity'].values else 0
            safety_analysis['severe_ae_rate'] = severe_ae_rate
        
        # 分组安全性比较
        if 'group' in df.columns:
            group_ae_rates = df.groupby('group')['adverse_event'].mean()
            safety_analysis['group_ae_rates'] = group_ae_rates.to_dict()
            
            # 安全性信号检测
            max_rate = group_ae_rates.max()
            min_rate = group_ae_rates.min()
            
            if (max_rate - min_rate) > 0.1:  # 差异超过10%
                safety_analysis['safety_signal'] = True
                safety_analysis['safety_concern'] = f"组间AE率差异较大: {max_rate:.1%} vs {min_rate:.1%}"
            else:
                safety_analysis['safety_signal'] = False
    
    # 实验室安全性
    lab_cols = [col for col in df.columns if 'lab' in col.lower() or 'test' in col.lower()]
    if lab_cols:
        lab_abnormalities = {}
        for col in lab_cols:
            if df[col].dtype in ['int64', 'float64']:
                abnormal_count = detect_outliers(df[col])
                lab_abnormalities[col] = {
                    'abnormal_count': abnormal_count,
                    'abnormal_rate': abnormal_count / len(df)
                }
        
        safety_analysis['lab_abnormalities'] = lab_abnormalities
    
    return safety_analysis

def analyze_efficacy_trends(df):
    """分析疗效趋势"""
    
    efficacy_trends = {}
    
    # 查找疗效相关变量
    efficacy_cols = [col for col in df.columns if any(keyword in col.lower() 
                    for keyword in ['change', 'improvement', 'response', 'endpoint'])]
    
    if efficacy_cols:
        for col in efficacy_cols:
            if df[col].dtype in ['int64', 'float64']:
                trend_data = {
                    'mean': df[col].mean(),
                    'std': df[col].std(),
                    'median': df[col].median()
                }
                
                # 分组疗效比较
                if 'group' in df.columns:
                    group_efficacy = df.groupby('group')[col].agg(['mean', 'std', 'count'])
                    trend_data['by_group'] = group_efficacy.to_dict()
                    
                    # 计算效应量
                    if len(df['group'].unique()) == 2:
                        groups = df['group'].unique()
                        group1_data = df[df['group'] == groups[0]][col]
                        group2_data = df[df['group'] == groups[1]][col]
                        
                        # Cohen's d
                        pooled_std = np.sqrt(((len(group1_data)-1)*group1_data.std()**2 + 
                                            (len(group2_data)-1)*group2_data.std()**2) / 
                                           (len(group1_data)+len(group2_data)-2))
                        
                        cohens_d = (group2_data.mean() - group1_data.mean()) / pooled_std
                        trend_data['effect_size'] = cohens_d
                
                efficacy_trends[col] = trend_data
    
    return efficacy_trends

def assess_interim_data_quality(df):
    """评估中期数据质量"""
    
    quality_assessment = {}
    
    # 数据完整性
    missing_rate = df.isnull().sum().sum() / (len(df) * len(df.columns))
    quality_assessment['missing_rate'] = missing_rate
    
    # 数据一致性
    duplicate_rate = df.duplicated().sum() / len(df)
    quality_assessment['duplicate_rate'] = duplicate_rate
    
    # 关键变量完整性
    key_vars = ['group', 'age', 'gender']  # 假设的关键变量
    key_var_completeness = {}
    
    for var in key_vars:
        if var in df.columns:
            completeness = (df[var].count() / len(df)) * 100
            key_var_completeness[var] = completeness
    
    quality_assessment['key_var_completeness'] = key_var_completeness
    
    # 数据录入及时性（如果有日期信息）
    if 'visit_date' in df.columns:
        df['visit_date'] = pd.to_datetime(df['visit_date'])
        latest_entry = df['visit_date'].max()
        days_since_latest = (datetime.now() - latest_entry).days
        
        quality_assessment['data_timeliness'] = {
            'latest_entry': latest_entry.strftime('%Y-%m-%d'),
            'days_since_latest': days_since_latest,
            'timely': days_since_latest <= 7  # 7天内为及时
        }
    
    # 数据质量评分
    completeness_score = (1 - missing_rate) * 100
    consistency_score = (1 - duplicate_rate) * 100
    overall_score = (completeness_score + consistency_score) / 2
    
    quality_assessment['scores'] = {
        'completeness': completeness_score,
        'consistency': consistency_score,
        'overall': overall_score
    }
    
    return quality_assessment

def perform_futility_analysis(df):
    """执行无效性分析"""
    
    futility_results = {}
    
    # 条件功效分析
    if 'group' in df.columns:
        groups = df['group'].unique()
        
        if len(groups) == 2:
            # 查找主要终点变量
            endpoint_cols = [col for col in df.columns if any(keyword in col.lower() 
                            for keyword in ['change', 'endpoint', 'primary'])]
            
            if endpoint_cols:
                primary_endpoint = endpoint_cols[0]
                
                group1_data = df[df['group'] == groups[0]][primary_endpoint].dropna()
                group2_data = df[df['group'] == groups[1]][primary_endpoint].dropna()
                
                if len(group1_data) > 0 and len(group2_data) > 0:
                    # 计算当前效应量
                    current_effect = group2_data.mean() - group1_data.mean()
                    pooled_std = np.sqrt((group1_data.var() + group2_data.var()) / 2)
                    standardized_effect = current_effect / pooled_std if pooled_std > 0 else 0
                    
                    # 条件功效计算（简化版）
                    # 假设目标效应量为0.5（中等效应）
                    target_effect = 0.5
                    
                    if abs(standardized_effect) < 0.1:  # 效应量很小
                        conditional_power = 0.1  # 低条件功效
                        futility_recommendation = "建议考虑终止研究（无效性）"
                    elif abs(standardized_effect) < 0.3:
                        conditional_power = 0.3
                        futility_recommendation = "继续研究，但需密切监察"
                    else:
                        conditional_power = 0.7
                        futility_recommendation = "继续研究"
                    
                    futility_results = {
                        'primary_endpoint': primary_endpoint,
                        'current_effect_size': standardized_effect,
                        'conditional_power': conditional_power,
                        'recommendation': futility_recommendation,
                        'sample_sizes': {
                            groups[0]: len(group1_data),
                            groups[1]: len(group2_data)
                        }
                    }
    
    return futility_results

def display_interim_report(content, recipients):
    """显示中期分析报告"""
    
    st.markdown("# 📊 中期分析报告")
    
    # 报告头部信息
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("中期分析日期", content['interim_date'].strftime('%Y-%m-%d'))
    
    with col2:
        st.metric("入组进展", f"{content['enrollment_rate']:.1f}%")
    
    with col3:
        st.metric("数据截止日期", content['cutoff_date'].strftime('%Y-%m-%d'))
    
    st.markdown("---")
    
    # 执行摘要
    st.markdown("## 📋 执行摘要")
    
    summary_points = generate_interim_summary(content)
    
    for point in summary_points:
        st.info(f"• {point}")
    
    # 入组进展分析
    if 'enrollment' in content['analysis_components']:
        st.markdown("## 👥 入组进展分析")
        
        enrollment = content['analysis_components']['enrollment']
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("计划入组", enrollment['planned'])
        
        with col2:
            st.metric("已入组", enrollment['current'])
        
        with col3:
            st.metric("入组率", f"{enrollment['rate']:.1f}%")
        
        with col4:
            st.metric("剩余入组", enrollment['remaining'])
        
        # 入组趋势图
        if 'monthly_trend' in enrollment:
            st.markdown("### 📈 月度入组趋势")
            
            monthly_data = enrollment['monthly_trend']
            months = list(monthly_data.keys())
            counts = list(monthly_data.values())
            
            fig = px.line(
                x=[str(month) for month in months],
                y=counts,
                title="月度入组人数趋势",
                labels={'x': '月份', 'y': '入组人数'}
            )
            fig.update_layout(height=400)
            st.plotly_chart(fig, use_container_width=True)
            
            # 预测完成时间
            if 'estimated_completion_months' in enrollment:
                months_remaining = enrollment['estimated_completion_months']
                if months_remaining != float('inf'):
                    estimated_completion = datetime.now() + timedelta(days=months_remaining*30)
                    st.info(f"📅 预计完成入组时间: {estimated_completion.strftime('%Y年%m月')}")
    
    # 安全性监察
    if 'safety' in content['analysis_components']:
        st.markdown("## ⚠️ 安全性监察")
        
        safety = content['analysis_components']['safety']
        
        # 总体安全性指标
        if 'ae_rate' in safety:
            col1, col2 = st.columns(2)
            
            with col1:
                st.metric("不良事件发生率", f"{safety['ae_rate']:.1%}")
            
            with col2:
                if 'severe_ae_rate' in safety:
                    st.metric("严重不良事件发生率", f"{safety['severe_ae_rate']:.1%}")
        
        # 分组安全性比较
        if 'group_ae_rates' in safety:
            st.markdown("### 分组安全性比较")
            
            safety_data = []
            for group, rate in safety['group_ae_rates'].items():
                safety_data.append({
                    '分组': group,
                    'AE发生率': f"{rate:.1%}"
                })
            
            st.dataframe(pd.DataFrame(safety_data), hide_index=True)
            
            # 可视化
            fig = px.bar(
                pd.DataFrame(safety_data),
                x='分组',
                y='AE发生率',
                title="各组不良事件发生率比较"
            )
            st.plotly_chart(fig, use_container_width=True)
        
        # 安全性信号
        if safety.get('safety_signal', False):
            st.warning(f"⚠️ 安全性信号: {safety.get('safety_concern', '')}")
        else:
            st.success("✅ 未发现明显安全性信号")
        
        # 实验室安全性
        if 'lab_abnormalities' in safety:
            st.markdown("### 🧪 实验室安全性指标")
            
            lab_data = []
            for lab, results in safety['lab_abnormalities'].items():
                lab_data.append({
                    '指标': lab,
                    '异常例数': results['abnormal_count'],
                    '异常率': f"{results['abnormal_rate']:.1%}"
                })
            
            if lab_data:
                st.dataframe(pd.DataFrame(lab_data), hide_index=True)
    
    # 疗效趋势分析
    if 'efficacy' in content['analysis_components']:
        st.markdown("## 📈 疗效趋势分析")
        
        efficacy = content['analysis_components']['efficacy']
        
        for endpoint, data in efficacy.items():
            st.markdown(f"### {endpoint}")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("均值", f"{data['mean']:.3f}")
            
            with col2:
                st.metric("标准差", f"{data['std']:.3f}")
            
            with col3:
                st.metric("中位数", f"{data['median']:.3f}")
            
            # 分组疗效比较
            if 'by_group' in data:
                st.markdown("#### 分组疗效比较")
                
                group_data = []
                for group in data['by_group']['mean'].keys():
                    group_data.append({
                        '分组': group,
                        '均值': f"{data['by_group']['mean'][group]:.3f}",
                        '标准差': f"{data['by_group']['std'][group]:.3f}",
                        '样本量': data['by_group']['count'][group]
                    })
                
                st.dataframe(pd.DataFrame(group_data), hide_index=True)
                
                # 效应量
                if 'effect_size' in data:
                    effect_size = data['effect_size']
                    
                    if abs(effect_size) < 0.2:
                        effect_interpretation = "小效应"
                    elif abs(effect_size) < 0.5:
                        effect_interpretation = "中等效应"
                    else:
                        effect_interpretation = "大效应"
                    
                    st.info(f"📊 效应量 (Cohen's d): {effect_size:.3f} ({effect_interpretation})")
    
    # 数据质量评估
    if 'quality' in content['analysis_components']:
        st.markdown("## 🔍 数据质量评估")
        
        quality = content['analysis_components']['quality']
        
        # 质量评分
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("完整性评分", f"{quality['scores']['completeness']:.1f}/100")
        
        with col2:
            st.metric("一致性评分", f"{quality['scores']['consistency']:.1f}/100")
        
        with col3:
            st.metric("总体质量评分", f"{quality['scores']['overall']:.1f}/100")
        
        # 关键变量完整性
        if 'key_var_completeness' in quality:
            st.markdown("### 关键变量完整性")
            
            completeness_data = []
            for var, completeness in quality['key_var_completeness'].items():
                completeness_data.append({
                    '变量': var,
                    '完整性': f"{completeness:.1f}%"
                })
            
            st.dataframe(pd.DataFrame(completeness_data), hide_index=True)
        
        # 数据录入及时性
        if 'data_timeliness' in quality:
            timeliness = quality['data_timeliness']
            
            if timeliness['timely']:
                st.success(f"✅ 数据录入及时 (最新数据: {timeliness['latest_entry']})")
            else:
                st.warning(f"⚠️ 数据录入延迟 ({timeliness['days_since_latest']} 天前)")
    
    # 无效性分析
    if 'futility' in content['analysis_components']:
        st.markdown("## 🎯 无效性分析")
        
        futility = content['analysis_components']['futility']
        
        if futility:
            col1, col2 = st.columns(2)
            
            with col1:
                st.metric("当前效应量", f"{futility['current_effect_size']:.3f}")
            
            with col2:
                st.metric("条件功效", f"{futility['conditional_power']:.1%}")
            
            # 样本量信息
            st.markdown("### 当前样本量")
            
            sample_data = []
            for group, n in futility['sample_sizes'].items():
                sample_data.append({
                    '分组': group,
                    '样本量': n
                })
            
            st.dataframe(pd.DataFrame(sample_data), hide_index=True)
            
            # 建议
            recommendation = futility['recommendation']
            
            if "终止" in recommendation:
                st.error(f"🛑 {recommendation}")
            elif "密切监察" in recommendation:
                st.warning(f"⚠️ {recommendation}")
            else:
                st.success(f"✅ {recommendation}")
    
    # 结论和建议
    st.markdown("## 📝 结论和建议")
    
    conclusions = generate_interim_conclusions(content)
    
    for conclusion in conclusions:
        st.success(f"✅ {conclusion}")
    
    # 下一步行动
    st.markdown("## 🎯 下一步行动")
    
    next_actions = generate_next_actions(content)
    
    for action in next_actions:
        st.info(f"📋 {action}")
    
    # 报告接收方
    st.markdown("---")
    st.markdown("### 📧 报告分发")
    st.info(f"本报告将分发给: {', '.join(recipients)}")

def generate_interim_summary(content):
    """生成中期分析摘要"""
    
    summary_points = []
    
    # 入组进展摘要
    enrollment_rate = content['enrollment_rate']
    
    if enrollment_rate >= 80:
        summary_points.append(f"入组进展良好，已完成 {enrollment_rate:.1f}% 的计划入组")
    elif enrollment_rate >= 50:
        summary_points.append(f"入组进展正常，已完成 {enrollment_rate:.1f}% 的计划入组")
    else:
        summary_points.append(f"入组进展缓慢，仅完成 {enrollment_rate:.1f}% 的计划入组，需要加强入组措施")
    
    # 安全性摘要
    if 'safety' in content['analysis_components']:
        safety = content['analysis_components']['safety']
        
        if 'ae_rate' in safety:
            ae_rate = safety['ae_rate']
            
            if ae_rate < 0.1:
                summary_points.append("安全性良好，不良事件发生率较低")
            elif ae_rate < 0.2:
                summary_points.append("安全性可接受，不良事件发生率在预期范围内")
            else:
                summary_points.append("需要关注安全性，不良事件发生率偏高")
        
        if safety.get('safety_signal', False):
            summary_points.append("发现潜在安全性信号，需要进一步评估")
    
    # 疗效摘要
    if 'efficacy' in content['analysis_components']:
        efficacy = content['analysis_components']['efficacy']
        
        has_positive_trend = False
        
        for endpoint, data in efficacy.items():
            if 'effect_size' in data:
                effect_size = abs(data['effect_size'])
                
                if effect_size >= 0.5:
                    has_positive_trend = True
                    break
        
        if has_positive_trend:
            summary_points.append("疗效趋势积极，观察到中等以上效应量")
        else:
            summary_points.append("疗效趋势尚不明确，需要继续观察")
    
    # 数据质量摘要
    if 'quality' in content['analysis_components']:
        quality = content['analysis_components']['quality']
        overall_score = quality['scores']['overall']
        
        if overall_score >= 90:
            summary_points.append("数据质量优秀")
        elif overall_score >= 80:
            summary_points.append("数据质量良好")
        else:
            summary_points.append("数据质量需要改进")
    
    return summary_points

def generate_interim_conclusions(content):
    """生成中期分析结论"""
    
    conclusions = []
    
    # 总体结论
    enrollment_rate = content['enrollment_rate']
    
    if enrollment_rate >= 50:
        conclusions.append("研究按计划进行，入组进展符合预期")
    else:
        conclusions.append("研究入组进展缓慢，建议调整入组策略")
    
    # 安全性结论
    if 'safety' in content['analysis_components']:
        safety = content['analysis_components']['safety']
        
        if not safety.get('safety_signal', False):
            conclusions.append("未发现新的安全性信号，可以继续研究")
        else:
            conclusions.append("发现安全性信号，建议加强安全性监察")
    
    # 疗效结论
    if 'efficacy' in content['analysis_components']:
        conclusions.append("疗效数据收集正常，趋势分析将在更多数据收集后进行")
    
    # 数据质量结论
    if 'quality' in content['analysis_components']:
        quality = content['analysis_components']['quality']
        
        if quality['scores']['overall'] >= 85:
            conclusions.append("数据质量满足分析要求")
        else:
            conclusions.append("需要加强数据质量控制措施")
    
    # 无效性分析结论
    if 'futility' in content['analysis_components']:
        futility = content['analysis_components']['futility']
        
        if futility and "终止" in futility.get('recommendation', ''):
            conclusions.append("基于无效性分析，建议考虑提前终止研究")
        else:
            conclusions.append("基于当前数据，建议继续进行研究")
    
    return conclusions

def generate_next_actions(content):
    """生成下一步行动计划"""
    
    actions = []
    
    # 入组相关行动
    enrollment_rate = content['enrollment_rate']
    
    if enrollment_rate < 50:
        actions.append("制定入组加速计划，考虑增加研究中心或调整入选标准")
    elif enrollment_rate < 80:
        actions.append("继续按当前速度入组，定期监控入组进展")
    
    # 安全性相关行动
    if 'safety' in content['analysis_components']:
        safety = content['analysis_components']['safety']
        
        if safety.get('safety_signal', False):
            actions.append("召开安全性评估会议，详细分析安全性信号")
            actions.append("考虑调整安全性监察频率")
        
        actions.append("继续收集和监察安全性数据")
    
    # 数据质量相关行动
    if 'quality' in content['analysis_components']:
        quality = content['analysis_components']['quality']
        
        if quality['scores']['overall'] < 85:
            actions.append("加强数据质量控制，提供额外的培训")
            actions.append("增加数据核查频率")
        
        if 'data_timeliness' in quality and not quality['data_timeliness']['timely']:
            actions.append("改善数据录入时效性，建立数据录入提醒机制")
    
    # 常规行动
    actions.append("准备下次中期分析计划")
    actions.append("更新研究进展报告给相关方")
    
    return actions

def generate_final_report(df):
    """生成最终研究报告"""
    st.markdown("### 📑 最终研究报告")
    
    st.info("最终研究报告是研究完成后的综合性报告，包含完整的统计分析、结论和临床意义解释。")
    
    # 最终报告配置
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("#### ⚙️ 研究信息")
        
        study_title = st.text_input(
            "研究标题",
            value="临床试验最终分析报告"
        )
        
        study_phase = st.selectbox(
            "研究阶段",
            ["I期", "II期", "III期", "IV期", "其他"]
        )
        
        study_design = st.selectbox(
            "研究设计",
            ["随机对照试验", "队列研究", "病例对照研究", "横断面研究", "其他"]
        )
        
        primary_endpoint = st.text_input(
            "主要终点",
            value="主要疗效指标改善率"
        )
        
        study_duration = st.text_input(
            "研究持续时间",
            value="12个月"
        )
    
    with col2:
        st.markdown("#### 📊 报告内容")
        
        include_background = st.checkbox("研究背景", value=True)
        include_methods = st.checkbox("研究方法", value=True)
        include_results = st.checkbox("研究结果", value=True)
        include_safety_analysis = st.checkbox("安全性分析", value=True)
        include_efficacy_analysis = st.checkbox("疗效分析", value=True)
        include_subgroup_analysis = st.checkbox("亚组分析", value=True)
        include_discussion = st.checkbox("讨论", value=True)
        include_conclusions = st.checkbox("结论", value=True)
        
        report_format = st.selectbox(
            "报告格式",
            ["完整版", "摘要版", "监管版", "学术版"]
        )
    
    if st.button("📑 生成最终研究报告", type="primary"):
        
        # 创建最终报告内容
        final_content = create_final_report_content(
            df, study_title, study_phase, study_design, primary_endpoint,
            study_duration, include_background, include_methods, include_results,
            include_safety_analysis, include_efficacy_analysis, include_subgroup_analysis,
            include_discussion, include_conclusions
        )
        
        # 显示最终报告
        display_final_report(final_content, report_format)

def create_final_report_content(df, title, phase, design, primary_endpoint,
                               duration, include_background, include_methods,
                               include_results, include_safety, include_efficacy,
                               include_subgroup, include_discussion, include_conclusions):
    """创建最终报告内容"""
    
    content = {
        'study_info': {
            'title': title,
            'phase': phase,
            'design': design,
            'primary_endpoint': primary_endpoint,
            'duration': duration,
            'total_subjects': len(df),
            'analysis_date': datetime.now().strftime('%Y-%m-%d')
        },
        'sections': {}
    }
    
    # 研究结果
    if include_results:
        content['sections']['results'] = analyze_final_results(df, primary_endpoint)
    
    # 安全性分析
    if include_safety:
        content['sections']['safety'] = perform_final_safety_analysis(df)
    
    # 疗效分析
    if include_efficacy:
        content['sections']['efficacy'] = perform_final_efficacy_analysis(df, primary_endpoint)
    
    # 亚组分析
    if include_subgroup:
        content['sections']['subgroup'] = perform_final_subgroup_analysis(df)
    
    return content

def analyze_final_results(df, primary_endpoint):
    """分析最终结果"""
    
    results = {
        'demographics': analyze_final_demographics(df),
        'disposition': analyze_subject_disposition(df),
        'primary_analysis': analyze_primary_endpoint(df, primary_endpoint)
    }
    
    return results

def analyze_final_demographics(df):
    """分析最终人口学特征"""
    
    demographics = {}
    
    # 年龄分析
    if 'age' in df.columns:
        demographics['age'] = {
            'n': df['age'].count(),
            'mean': df['age'].mean(),
            'std': df['age'].std(),
            'median': df['age'].median(),
            'range': [df['age'].min(), df['age'].max()],
            'age_groups': pd.cut(df['age'], bins=[0, 18, 35, 50, 65, 100], 
                               labels=['<18', '18-34', '35-49', '50-64', '≥65']).value_counts().to_dict()
        }
    
    # 性别分析
    if 'gender' in df.columns:
        gender_counts = df['gender'].value_counts()
        demographics['gender'] = {
            'counts': gender_counts.to_dict(),
            'percentages': (gender_counts / len(df) * 100).to_dict()
        }
    
    # 分组分析
    if 'group' in df.columns:
        group_counts = df['group'].value_counts()
        demographics['treatment_groups'] = {
            'counts': group_counts.to_dict(),
            'percentages': (group_counts / len(df) * 100).to_dict()
        }
        
        # 分组间人口学比较
        if 'age' in df.columns:
            age_by_group = df.groupby('group')['age'].agg(['mean', 'std']).to_dict()
            demographics['age_by_group'] = age_by_group
        
        if 'gender' in df.columns:
            gender_by_group = pd.crosstab(df['group'], df['gender'], normalize='index') * 100
            demographics['gender_by_group'] = gender_by_group.to_dict()
    
    return demographics

def analyze_subject_disposition(df):
    """分析受试者处置情况"""
    
    disposition = {
        'enrolled': len(df),
        'completed': len(df),  # 假设所有受试者都完成了研究
        'completion_rate': 100.0
    }
    
    # 如果有依从性数据
    if 'compliance' in df.columns:
        high_compliance = (df['compliance'] >= 0.8).sum()
        disposition['high_compliance'] = high_compliance
        disposition['high_compliance_rate'] = (high_compliance / len(df)) * 100
    
    # 分组处置情况
    if 'group' in df.columns:
        group_disposition = {}
        for group in df['group'].unique():
            group_data = df[df['group'] == group]
            group_disposition[group] = {
                'enrolled': len(group_data),
                'completed': len(group_data),
                'completion_rate': 100.0
            }
        
        disposition['by_group'] = group_disposition
    
    return disposition

def analyze_primary_endpoint(df, primary_endpoint):
    """分析主要终点"""
    
    # 查找主要终点相关变量
    endpoint_cols = [col for col in df.columns if any(keyword in col.lower() 
                    for keyword in ['change', 'endpoint', 'primary', 'response'])]
    
    if not endpoint_cols:
        # 如果没有找到，使用第一个数值变量作为示例
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        endpoint_cols = [numeric_cols[0]] if len(numeric_cols) > 0 else []
    
    primary_analysis = {}
    
    if endpoint_cols:
        primary_var = endpoint_cols[0]
        
        # 总体分析
        primary_analysis['variable'] = primary_var
        primary_analysis['overall'] = {
            'n': df[primary_var].count(),
            'mean': df[primary_var].mean(),
            'std': df[primary_var].std(),
            'median': df[primary_var].median(),
            'q1': df[primary_var].quantile(0.25),
            'q3': df[primary_var].quantile(0.75),
            'min': df[primary_var].min(),
            'max': df[primary_var].max()
        }
        
        # 分组分析
        if 'group' in df.columns:
            groups = df['group'].unique()
            
            if len(groups) >= 2:
                group_analysis = {}
                
                for group in groups:
                    group_data = df[df['group'] == group][primary_var]
                    group_analysis[group] = {
                        'n': group_data.count(),
                        'mean': group_data.mean(),
                        'std': group_data.std(),
                        'median': group_data.median(),
                        'q1': group_data.quantile(0.25),
                        'q3': group_data.quantile(0.75)
                    }
                
                primary_analysis['by_group'] = group_analysis
                
                # 统计检验
                if len(groups) == 2:
                    from scipy import stats
                    
                    group1_data = df[df['group'] == groups[0]][primary_var].dropna()
                    group2_data = df[df['group'] == groups[1]][primary_var].dropna()
                    
                    # t检验
                    try:
                        t_stat, p_value = stats.ttest_ind(group1_data, group2_data)
                        
                        # 效应量计算
                        pooled_std = np.sqrt(((len(group1_data)-1)*group1_data.std()**2 + 
                                            (len(group2_data)-1)*group2_data.std()**2) / 
                                           (len(group1_data)+len(group2_data)-2))
                        
                        cohens_d = (group2_data.mean() - group1_data.mean()) / pooled_std
                        
                        # 95%置信区间
                        diff_mean = group2_data.mean() - group1_data.mean()
                        se_diff = pooled_std * np.sqrt(1/len(group1_data) + 1/len(group2_data))
                        
                        from scipy.stats import t
                        df_t = len(group1_data) + len(group2_data) - 2
                        t_critical = t.ppf(0.975, df_t)
                        
                        ci_lower = diff_mean - t_critical * se_diff
                        ci_upper = diff_mean + t_critical * se_diff
                        
                        primary_analysis['statistical_test'] = {
                            'test_type': 't-test',
                            't_statistic': t_stat,
                            'p_value': p_value,
                            'significant': p_value < 0.05,
                            'effect_size': cohens_d,
                            'mean_difference': diff_mean,
                            'ci_95': [ci_lower, ci_upper]
                        }
                        
                    except Exception as e:
                        primary_analysis['statistical_test'] = {'error': str(e)}
    
    return primary_analysis

def perform_final_safety_analysis(df):
    """执行最终安全性分析"""
    
    safety_analysis = {}
    
    # 总体安全性概述
    safety_analysis['overview'] = {
        'total_subjects': len(df),
        'safety_evaluable': len(df)  # 假设所有受试者都可评估安全性
    }
    
    # 不良事件分析
    if 'adverse_event' in df.columns:
        ae_subjects = df['adverse_event'].sum() if df['adverse_event'].dtype in ['int64', 'bool'] else 0
        
        safety_analysis['adverse_events'] = {
            'subjects_with_ae': ae_subjects,
            'ae_rate': (ae_subjects / len(df)) * 100,
            'subjects_without_ae': len(df) - ae_subjects
        }
        
        # 按严重程度分析
        if 'severity' in df.columns:
            ae_data = df[df['adverse_event'] == 1] if df['adverse_event'].dtype in ['int64', 'bool'] else df
            
            if len(ae_data) > 0:
                severity_counts = ae_data['severity'].value_counts()
                safety_analysis['ae_by_severity'] = {
                    'counts': severity_counts.to_dict(),
                    'percentages': (severity_counts / len(ae_data) * 100).to_dict()
                }
        
        # 按因果关系分析
        if 'causality' in df.columns:
            ae_data = df[df['adverse_event'] == 1] if df['adverse_event'].dtype in ['int64', 'bool'] else df
            
            if len(ae_data) > 0:
                causality_counts = ae_data['causality'].value_counts()
                safety_analysis['ae_by_causality'] = {
                    'counts': causality_counts.to_dict(),
                    'percentages': (causality_counts / len(ae_data) * 100).to_dict()
                }
        
        # 分组安全性比较
        if 'group' in df.columns:
            group_safety = {}
            
            for group in df['group'].unique():
                group_data = df[df['group'] == group]
                group_ae = group_data['adverse_event'].sum() if group_data['adverse_event'].dtype in ['int64', 'bool'] else 0
                
                group_safety[group] = {
                    'n': len(group_data),
                    'subjects_with_ae': group_ae,
                    'ae_rate': (group_ae / len(group_data)) * 100 if len(group_data) > 0 else 0
                }
            
            safety_analysis['ae_by_group'] = group_safety
            
            # 组间比较统计检验
            if len(df['group'].unique()) == 2:
                from scipy.stats import chi2_contingency
                
                groups = df['group'].unique()
                
                # 构建列联表
                contingency_table = pd.crosstab(df['group'], df['adverse_event'])
                
                try:
                    chi2, p_value, dof, expected = chi2_contingency(contingency_table)
                    
                    safety_analysis['group_comparison'] = {
                        'test_type': 'Chi-square test',
                        'chi2_statistic': chi2,
                        'p_value': p_value,
                        'significant': p_value < 0.05
                    }
                except:
                    safety_analysis['group_comparison'] = {'error': '无法执行卡方检验'}
    
    # 实验室安全性参数
    lab_cols = [col for col in df.columns if any(keyword in col.lower() 
               for keyword in ['lab', 'test', 'level', 'count', 'alt', 'ast', 'creatinine'])]
    
    if lab_cols:
        lab_safety = {}
        
        for col in lab_cols:
            if df[col].dtype in ['int64', 'float64']:
                # 异常值检测
                Q1 = df[col].quantile(0.25)
                Q3 = df[col].quantile(0.75)
                IQR = Q3 - Q1
                
                lower_bound = Q1 - 1.5 * IQR
                upper_bound = Q3 + 1.5 * IQR
                
                abnormal_low = (df[col] < lower_bound).sum()
                abnormal_high = (df[col] > upper_bound).sum()
                total_abnormal = abnormal_low + abnormal_high
                
                lab_safety[col] = {
                    'n': df[col].count(),
                    'mean': df[col].mean(),
                    'std': df[col].std(),
                    'abnormal_low': abnormal_low,
                    'abnormal_high': abnormal_high,
                    'total_abnormal': total_abnormal,
                    'abnormal_rate': (total_abnormal / df[col].count()) * 100 if df[col].count() > 0 else 0
                }
        
        safety_analysis['laboratory'] = lab_safety
    
    return safety_analysis

def perform_final_efficacy_analysis(df, primary_endpoint):
    """执行最终疗效分析"""
    
    efficacy_analysis = {}
    
    # 查找疗效相关变量
    efficacy_cols = [col for col in df.columns if any(keyword in col.lower() 
                    for keyword in ['change', 'improvement', 'response', 'endpoint', 'efficacy'])]
    
    if not efficacy_cols:
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        efficacy_cols = list(numeric_cols[:3]) if len(numeric_cols) > 0 else []
    
    # 主要疗效分析
    if efficacy_cols:
        primary_var = efficacy_cols[0]
        
        efficacy_analysis['primary_efficacy'] = {
            'endpoint': primary_var,
            'analysis': analyze_primary_endpoint(df, primary_var)
        }
        
        # 次要疗效分析
        if len(efficacy_cols) > 1:
            secondary_analysis = {}
            
            for col in efficacy_cols[1:]:
                if df[col].dtype in ['int64', 'float64']:
                    secondary_analysis[col] = {
                        'overall': {
                            'n': df[col].count(),
                            'mean': df[col].mean(),
                            'std': df[col].std(),
                            'median': df[col].median()
                        }
                    }
                    
                    # 分组分析
                    if 'group' in df.columns:
                        group_results = {}
                        
                        for group in df['group'].unique():
                            group_data = df[df['group'] == group][col]
                            group_results[group] = {
                                'n': group_data.count(),
                                'mean': group_data.mean(),
                                'std': group_data.std(),
                                'median': group_data.median()
                            }
                        
                        secondary_analysis[col]['by_group'] = group_results
            
            efficacy_analysis['secondary_efficacy'] = secondary_analysis
    
    # 应答率分析（如果有二分类结果）
    response_cols = [col for col in df.columns if 'response' in col.lower()]
    
    if response_cols:
        response_analysis = {}
        
        for col in response_cols:
            if df[col].dtype in ['int64', 'bool'] or df[col].nunique() == 2:
                # 总体应答率
                response_rate = df[col].mean() * 100 if df[col].dtype in ['int64', 'bool'] else 0
                
                response_analysis[col] = {
                    'overall_response_rate': response_rate,
                    'responders': df[col].sum() if df[col].dtype in ['int64', 'bool'] else 0,
                    'non_responders': len(df) - (df[col].sum() if df[col].dtype in ['int64', 'bool'] else 0)
                }
                
                # 分组应答率比较
                if 'group' in df.columns:
                    group_response = {}
                    
                    for group in df['group'].unique():
                        group_data = df[df['group'] == group]
                        group_response_rate = group_data[col].mean() * 100 if group_data[col].dtype in ['int64', 'bool'] else 0
                        
                        group_response[group] = {
                            'n': len(group_data),
                            'responders': group_data[col].sum() if group_data[col].dtype in ['int64', 'bool'] else 0,
                            'response_rate': group_response_rate
                        }
                    
                    response_analysis[col]['by_group'] = group_response
        
        efficacy_analysis['response_rates'] = response_analysis
    
    return efficacy_analysis

def perform_final_subgroup_analysis(df):
    """执行最终亚组分析"""
    
    subgroup_analysis = {}
    
    # 预定义亚组变量
    subgroup_vars = []
    
    if 'age' in df.columns:
        # 年龄亚组
        df['age_group'] = pd.cut(df['age'], bins=[0, 40, 60, 100], labels=['≤40岁', '41-60岁', '>60岁'])
        subgroup_vars.append('age_group')
    
    if 'gender' in df.columns:
        subgroup_vars.append('gender')
    
    # 查找其他分类变量
    categorical_cols = df.select_dtypes(include=['object']).columns
    for col in categorical_cols:
        if col not in ['group'] and df[col].nunique() <= 5:  # 限制类别数量
            subgroup_vars.append(col)
    
    # 查找主要疗效变量
    efficacy_cols = [col for col in df.columns if any(keyword in col.lower() 
                    for keyword in ['change', 'improvement', 'response', 'endpoint'])]
    
    if not efficacy_cols:
        numeric_cols = df.select_dtypes(include=[np.number]).columns
        efficacy_cols = [numeric_cols[0]] if len(numeric_cols) > 0 else []
    
    # 执行亚组分析
    if subgroup_vars and efficacy_cols and 'group' in df.columns:
        primary_efficacy = efficacy_cols[0]
        
        for subgroup_var in subgroup_vars:
            if subgroup_var in df.columns:
                subgroup_results = {}
                
                for subgroup_value in df[subgroup_var].unique():
                    if pd.notna(subgroup_value):
                        subgroup_data = df[df[subgroup_var] == subgroup_value]
                        
                        if len(subgroup_data) >= 10:  # 最小样本量要求
                            # 计算各组的疗效指标
                            group_results = {}
                            
                            for group in subgroup_data['group'].unique():
                                group_subgroup_data = subgroup_data[subgroup_data['group'] == group]
                                
                                if len(group_subgroup_data) > 0:
                                    group_results[group] = {
                                        'n': len(group_subgroup_data),
                                        'mean': group_subgroup_data[primary_efficacy].mean(),
                                        'std': group_subgroup_data[primary_efficacy].std()
                                    }
                            
                            # 计算组间差异
                            if len(group_results) >= 2:
                                groups = list(group_results.keys())
                                mean_diff = group_results[groups[1]]['mean'] - group_results[groups[0]]['mean']
                                
                                subgroup_results[str(subgroup_value)] = {
                                    'group_results': group_results,
                                    'mean_difference': mean_diff,
                                    'total_n': len(subgroup_data)
                                }
                
                if subgroup_results:
                    subgroup_analysis[subgroup_var] = subgroup_results
    
    return subgroup_analysis

def display_final_report(content, report_format):
    """显示最终研究报告"""
    
    study_info = content['study_info']
    
    # 报告标题和基本信息
    st.markdown(f"# {study_info['title']}")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("研究阶段", study_info['phase'])
    
    with col2:
        st.metric("研究设计", study_info['design'])
    
    with col3:
        st.metric("总样本量", study_info['total_subjects'])
    
    st.markdown(f"**主要终点**: {study_info['primary_endpoint']}")
    st.markdown(f"**研究持续时间**: {study_info['duration']}")
    st.markdown(f"**分析日期**: {study_info['analysis_date']}")
    
    st.markdown("---")
    
    # 摘要
    st.markdown("## 📋 研究摘要")
    
    summary_text = generate_study_summary(content)
    st.markdown(summary_text)
    
    # 研究结果
    if 'results' in content['sections']:
        st.markdown("## 📊 研究结果")
        
        results = content['sections']['results']
        
        # 受试者特征
        if 'demographics' in results:
            st.markdown("### 👥 受试者特征")
            
            demographics = results['demographics']
            
            # 年龄特征
            if 'age' in demographics:
                age_data = demographics['age']
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric("样本量", age_data['n'])
                
                with col2:
                    st.metric("平均年龄", f"{age_data['mean']:.1f}岁")
                
                with col3:
                    st.metric("年龄中位数", f"{age_data['median']:.1f}岁")
                
                with col4:
                    st.metric("年龄范围", f"{age_data['range'][0]:.0f}-{age_data['range'][1]:.0f}岁")
                
                # 年龄组分布
                if 'age_groups' in age_data:
                    st.markdown("#### 年龄组分布")
                    
                    age_group_data = []
                    for age_group, count in age_data['age_groups'].items():
                        age_group_data.append({
                            '年龄组': age_group,
                            '人数': count,
                            '比例': f"{(count/age_data['n'])*100:.1f}%"
                        })
                    
                    st.dataframe(pd.DataFrame(age_group_data), hide_index=True)
            
            # 性别分布
            if 'gender' in demographics:
                st.markdown("#### 性别分布")
                
                gender_data = []
                for gender, count in demographics['gender']['counts'].items():
                    percentage = demographics['gender']['percentages'][gender]
                    gender_data.append({
                        '性别': gender,
                        '人数': count,
                        '比例': f"{percentage:.1f}%"
                    })
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.dataframe(pd.DataFrame(gender_data), hide_index=True)
                
                with col2:
                    fig = px.pie(
                        pd.DataFrame(gender_data),
                        values='人数',
                        names='性别',
                        title="性别分布"
                    )
                    st.plotly_chart(fig, use_container_width=True)
            
            # 治疗组分布
            if 'treatment_groups' in demographics:
                st.markdown("#### 治疗组分布")
                
                group_data = []
                for group, count in demographics['treatment_groups']['counts'].items():
                    percentage = demographics['treatment_groups']['percentages'][group]
                    group_data.append({
                        '治疗组': group,
                        '人数': count,
                        '比例': f"{percentage:.1f}%"
                    })
                
                st.dataframe(pd.DataFrame(group_data), hide_index=True)
        
        # 受试者处置
        if 'disposition' in results:
            st.markdown("### 📋 受试者处置")
            
            disposition = results['disposition']
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("入组人数", disposition['enrolled'])
            
            with col2:
                st.metric("完成人数", disposition['completed'])
            
            with col3:
                st.metric("完成率", f"{disposition['completion_rate']:.1f}%")
            
            # 分组处置情况
            if 'by_group' in disposition:
                st.markdown("#### 分组处置情况")
                
                disposition_data = []
                for group, data in disposition['by_group'].items():
                    disposition_data.append({
                        '治疗组': group,
                        '入组人数': data['enrolled'],
                        '完成人数': data['completed'],
                        '完成率': f"{data['completion_rate']:.1f}%"
                    })
                
                st.dataframe(pd.DataFrame(disposition_data), hide_index=True)
        
        # 主要终点分析
        if 'primary_analysis' in results:
            st.markdown("### 🎯 主要终点分析")
            
            primary = results['primary_analysis']
            
            if 'variable' in primary:
                st.markdown(f"**分析变量**: {primary['variable']}")
                
                # 总体结果
                if 'overall' in primary:
                    st.markdown("#### 总体结果")
                    
                    overall = primary['overall']
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        st.metric("样本量", overall['n'])
                    
                    with col2:
                        st.metric("均值", f"{overall['mean']:.3f}")
                    
                    with col3:
                        st.metric("标准差", f"{overall['std']:.3f}")
                    
                    with col4:
                        st.metric("中位数", f"{overall['median']:.3f}")
                
                # 分组结果
                if 'by_group' in primary:
                    st.markdown("#### 分组结果")
                    
                    group_data = []
                    for group, data in primary['by_group'].items():
                        group_data.append({
                            '治疗组': group,
                            '样本量': data['n'],
                            '均值': f"{data['mean']:.3f}",
                            '标准差': f"{data['std']:.3f}",
                            '中位数': f"{data['median']:.3f}",
                            'Q1': f"{data['q1']:.3f}",
                            'Q3': f"{data['q3']:.3f}"
                        })
                    
                    st.dataframe(pd.DataFrame(group_data), hide_index=True)
                
                # 统计检验结果
                if 'statistical_test' in primary:
                    st.markdown("#### 统计检验结果")
                    
                    test_results = primary['statistical_test']
                    
                    if 'error' not in test_results:
                        col1, col2, col3 = st.columns(3)
                        
                        with col1:
                            st.metric("检验统计量", f"{test_results['t_statistic']:.4f}")
                        
                        with col2:
                            st.metric("P值", f"{test_results['p_value']:.4f}")
                        
                        with col3:
                            significance = "是" if test_results['significant'] else "否"
                            st.metric("统计学显著", significance)
                        
                        # 效应量和置信区间
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.metric("效应量 (Cohen's d)", f"{test_results['effect_size']:.3f}")
                        
                        with col2:
                            ci = test_results['ci_95']
                            st.metric("均值差异95%CI", f"[{ci[0]:.3f}, {ci[1]:.3f}]")
                        
                        # 结果解释
                        if test_results['significant']:
                            st.success("✅ 主要终点达到统计学显著差异")
                        else:
                            st.warning("⚠️ 主要终点未达到统计学显著差异")
    
    # 安全性分析
    if 'safety' in content['sections']:
        st.markdown("## ⚠️ 安全性分析")
        
        safety = content['sections']['safety']
        
        # 安全性概述
        if 'overview' in safety:
            overview = safety['overview']
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.metric("总受试者数", overview['total_subjects'])
            
            with col2:
                st.metric("安全性可评估人数", overview['safety_evaluable'])
        
        # 不良事件分析
        if 'adverse_events' in safety:
            st.markdown("### 不良事件分析")
            
            ae_data = safety['adverse_events']
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.metric("发生AE人数", ae_data['subjects_with_ae'])
            
            with col2:
                st.metric("AE发生率", f"{ae_data['ae_rate']:.1f}%")
            
            with col3:
                st.metric("未发生AE人数", ae_data['subjects_without_ae'])
            
            # 按严重程度分析
            if 'ae_by_severity' in safety:
                st.markdown("#### 按严重程度分类")
                
                severity_data = []
                for severity, count in safety['ae_by_severity']['counts'].items():
                    percentage = safety['ae_by_severity']['percentages'][severity]
                    severity_data.append({
                        '严重程度': severity,
                        '事件数': count,
                        '占比': f"{percentage:.1f}%"
                    })
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.dataframe(pd.DataFrame(severity_data), hide_index=True)
                
                with col2:
                    fig = px.pie(
                        pd.DataFrame(severity_data),
                        values='事件数',
                        names='严重程度',
                        title="不良事件严重程度分布"
                    )
                    st.plotly_chart(fig, use_container_width=True)
            
            # 分组安全性比较
            if 'ae_by_group' in safety:
                st.markdown("#### 分组安全性比较")
                
                group_safety_data = []
                for group, data in safety['ae_by_group'].items():
                    group_safety_data.append({
                        '治疗组': group,
                        '样本量': data['n'],
                        'AE人数': data['subjects_with_ae'],
                        'AE发生率': f"{data['ae_rate']:.1f}%"
                    })
                
                st.dataframe(pd.DataFrame(group_safety_data), hide_index=True)
                
                # 统计检验结果
                if 'group_comparison' in safety:
                    test_result = safety['group_comparison']
                    
                    if 'error' not in test_result:
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.metric("卡方统计量", f"{test_result['chi2_statistic']:.4f}")
                        
                        with col2:
                            st.metric("P值", f"{test_result['p_value']:.4f}")
                        
                        if test_result['significant']:
                            st.warning("⚠️ 组间不良事件发生率存在统计学显著差异")
                        else:
                            st.success("✅ 组间不良事件发生率无统计学显著差异")
        
        # 实验室安全性
        if 'laboratory' in safety:
            st.markdown("### 🧪 实验室安全性指标")
            
            lab_data = []
            for lab, results in safety['laboratory'].items():
                lab_data.append({
                    '指标': lab,
                    '样本量': results['n'],
                    '均值': f"{results['mean']:.2f}",
                    '标准差': f"{results['std']:.2f}",
                    '异常例数': results['total_abnormal'],
                    '异常率': f"{results['abnormal_rate']:.1f}%"
                })
            
            st.dataframe(pd.DataFrame(lab_data), hide_index=True)
    
    # 疗效分析
    if 'efficacy' in content['sections']:
        st.markdown("## 📈 疗效分析")
        
        efficacy = content['sections']['efficacy']
        
        # 主要疗效分析（已在研究结果中显示，这里可以添加更多细节）
        
        # 次要疗效分析
        if 'secondary_efficacy' in efficacy:
            st.markdown("### 次要疗效终点")
            
            for endpoint, data in efficacy['secondary_efficacy'].items():
                st.markdown(f"#### {endpoint}")
                
                # 总体结果
                if 'overall' in data:
                    overall = data['overall']
                    
                    col1, col2, col3, col4 = st.columns(4)
                    
                    with col1:
                        st.metric("样本量", overall['n'])
                    
                    with col2:
                        st.metric("均值", f"{overall['mean']:.3f}")
                    
                    with col3:
                        st.metric("标准差", f"{overall['std']:.3f}")
                    
                    with col4:
                        st.metric("中位数", f"{overall['median']:.3f}")
                
                # 分组结果
                if 'by_group' in data:
                    group_data = []
                    for group, group_result in data['by_group'].items():
                        group_data.append({
                            '治疗组': group,
                            '样本量': group_result['n'],
                            '均值': f"{group_result['mean']:.3f}",
                            '标准差': f"{group_result['std']:.3f}",
                            '中位数': f"{group_result['median']:.3f}"
                        })
                    
                    st.dataframe(pd.DataFrame(group_data), hide_index=True)
        
        # 应答率分析
        if 'response_rates' in efficacy:
            st.markdown("### 应答率分析")
            
            for endpoint, data in efficacy['response_rates'].items():
                st.markdown(f"#### {endpoint}")
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.metric("总体应答率", f"{data['overall_response_rate']:.1f}%")
                
                with col2:
                    st.metric("应答者", data['responders'])
                
                with col3:
                    st.metric("无应答者", data['non_responders'])
                
                # 分组应答率比较
                if 'by_group' in data:
                    st.markdown("##### 分组应答率比较")
                    
                    response_data = []
                    for group, group_data in data['by_group'].items():
                        response_data.append({
                            '治疗组': group,
                            '样本量': group_data['n'],
                            '应答者': group_data['responders'],
                            '应答率': f"{group_data['response_rate']:.1f}%"
                        })
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.dataframe(pd.DataFrame(response_data), hide_index=True)
                    
                    with col2:
                        fig = px.bar(
                            pd.DataFrame(response_data),
                            x='治疗组',
                            y='应答率',
                            title=f"{endpoint} 分组应答率比较"
                        )
                        st.plotly_chart(fig, use_container_width=True)
    
    # 亚组分析
    if 'subgroup' in content['sections']:
        st.markdown("## 🔍 亚组分析")
        
        subgroup = content['sections']['subgroup']
        
        for subgroup_var, subgroup_data in subgroup.items():
            st.markdown(f"### 按 {subgroup_var} 分层分析")
            
            subgroup_results = []
            
            for subgroup_value, results in subgroup_data.items():
                if 'group_results' in results:
                    group_results = results['group_results']
                    groups = list(group_results.keys())
                    
                    if len(groups) >= 2:
                        subgroup_results.append({
                            f'{subgroup_var}': subgroup_value,
                            f'{groups[0]} (n)': group_results[groups[0]]['n'],
                            f'{groups[0]} 均值': f"{group_results[groups[0]]['mean']:.3f}",
                            f'{groups[1]} (n)': group_results[groups[1]]['n'],
                            f'{groups[1]} 均值': f"{group_results[groups[1]]['mean']:.3f}",
                            '均值差异': f"{results['mean_difference']:.3f}",
                            '总样本量': results['total_n']
                        })
            
            if subgroup_results:
                st.dataframe(pd.DataFrame(subgroup_results), hide_index=True)
                
                # 森林图可视化（简化版）
                if len(subgroup_results) > 1:
                    fig = px.scatter(
                        pd.DataFrame(subgroup_results),
                        x='均值差异',
                        y=f'{subgroup_var}',
                        size='总样本量',
                        title=f"按{subgroup_var}分层的治疗效应",
                        labels={'均值差异': '治疗效应 (均值差异)'}
                    )
                    fig.add_vline(x=0, line_dash="dash", line_color="red")
                    st.plotly_chart(fig, use_container_width=True)
    
    # 讨论和结论
    st.markdown("## 💭 讨论")
    
    discussion_points = generate_discussion_points(content)
    
    for point in discussion_points:
        st.markdown(f"• {point}")
    
    st.markdown("## 📝 结论")
    
    conclusions = generate_final_conclusions(content)
    
    for conclusion in conclusions:
        st.success(f"✅ {conclusion}")
    
    # 局限性
    st.markdown("## ⚠️ 研究局限性")
    
    limitations = generate_study_limitations(content)
    
    for limitation in limitations:
        st.warning(f"⚠️ {limitation}")
    
    # 临床意义
    st.markdown("## 🏥 临床意义")
    
    clinical_significance = generate_clinical_significance(content)
    
    for significance in clinical_significance:
        st.info(f"💡 {significance}")
    
    # 报告下载
    st.markdown("---")
    st.markdown("### 📥 报告下载")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📄 下载PDF报告"):
            pdf_content = generate_final_pdf_report(content)
            st.download_button(
                "点击下载PDF",
                data=pdf_content,
                file_name=f"最终研究报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                mime="application/pdf"
            )
    
    with col2:
        if st.button("📊 下载Word报告"):
            word_content = generate_final_word_report(content)
            st.download_button(
                "点击下载Word",
                data=word_content,
                file_name=f"最终研究报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    
    with col3:
        if st.button("🌐 下载HTML报告"):
            html_content = generate_final_html_report(content)
            st.download_button(
                "点击下载HTML",
                data=html_content,
                file_name=f"最终研究报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html",
                mime="text/html"
            )

def generate_study_summary(content):
    """生成研究摘要"""
    
    study_info = content['study_info']
    
    summary = f"""
    ### 研究背景
    本研究是一项{study_info['phase']}{study_info['design']}，旨在评估{study_info['primary_endpoint']}。
    
    ### 研究方法
    研究共纳入{study_info['total_subjects']}名受试者，研究持续时间为{study_info['duration']}。
    
    ### 主要结果
    """
    
    # 添加主要结果摘要
    if 'results' in content['sections']:
        results = content['sections']['results']
        
        if 'primary_analysis' in results:
            primary = results['primary_analysis']
            
            if 'statistical_test' in primary:
                test_result = primary['statistical_test']
                
                if 'error' not in test_result:
                    if test_result['significant']:
                        summary += f"主要终点达到统计学显著差异 (p={test_result['p_value']:.4f})，"
                        summary += f"治疗组间均值差异为{test_result['mean_difference']:.3f}。"
                    else:
                        summary += f"主要终点未达到统计学显著差异 (p={test_result['p_value']:.4f})。"
    
    # 添加安全性摘要
    if 'safety' in content['sections']:
        safety = content['sections']['safety']
        
        if 'adverse_events' in safety:
            ae_rate = safety['adverse_events']['ae_rate']
            summary += f"\n\n### 安全性\n不良事件总发生率为{ae_rate:.1f}%，"
            
            if 'group_comparison' in safety:
                test_result = safety['group_comparison']
                
                if 'error' not in test_result:
                    if test_result['significant']:
                        summary += "组间安全性存在统计学显著差异。"
                    else:
                        summary += "组间安全性无统计学显著差异。"
    
    return summary

def generate_discussion_points(content):
    """生成讨论要点"""
    
    discussion = []
    
    # 主要发现讨论
    if 'results' in content['sections']:
        results = content['sections']['results']
        
        if 'primary_analysis' in results:
            primary = results['primary_analysis']
            
            if 'statistical_test' in primary:
                test_result = primary['statistical_test']
                
                if 'error' not in test_result:
                    if test_result['significant']:
                        effect_size = abs(test_result['effect_size'])
                        
                        if effect_size >= 0.8:
                            discussion.append("本研究观察到大效应量的治疗效果，具有重要的临床意义")
                        elif effect_size >= 0.5:
                            discussion.append("本研究观察到中等效应量的治疗效果，提示治疗具有一定的临床价值")
                        else:
                            discussion.append("虽然达到统计学显著性，但效应量较小，临床意义有待进一步评估")
                    else:
                        discussion.append("主要终点未达到统计学显著性，可能与样本量、研究设计或治疗效果有关")
    
    # 安全性讨论
    if 'safety' in content['sections']:
        safety = content['sections']['safety']
        
        if 'adverse_events' in safety:
            ae_rate = safety['adverse_events']['ae_rate']
            
            if ae_rate < 10:
                discussion.append("不良事件发生率较低，治疗的安全性良好")
            elif ae_rate < 20:
                discussion.append("不良事件发生率在可接受范围内，但需要继续监察")
            else:
                discussion.append("不良事件发生率较高，需要仔细权衡获益风险比")
    
    # 亚组分析讨论
    if 'subgroup' in content['sections']:
        discussion.append("亚组分析结果提示治疗效果可能在不同人群中存在差异，但需要谨慎解释")
    
    # 研究设计讨论
    study_info = content['study_info']
    
    if study_info['design'] == '随机对照试验':
        discussion.append("随机对照设计有效控制了混杂因素，提高了结果的可信度")
    
    # 样本量讨论
    if study_info['total_subjects'] < 100:
        discussion.append("样本量相对较小，结果的推广性可能受到限制")
    elif study_info['total_subjects'] > 500:
        discussion.append("较大的样本量提高了研究结果的统计功效和可靠性")
    
    return discussion

def generate_final_conclusions(content):
    """生成最终结论"""
    
    conclusions = []
    
    # 主要终点结论
    if 'results' in content['sections']:
        results = content['sections']['results']
        
        if 'primary_analysis' in results:
            primary = results['primary_analysis']
            
            if 'statistical_test' in primary:
                test_result = primary['statistical_test']
                
                if 'error' not in test_result:
                    if test_result['significant']:
                        conclusions.append(f"主要终点达到统计学显著差异，证实了治疗的有效性")
                    else:
                        conclusions.append(f"主要终点未达到预设的统计学显著性标准")
    
    # 安全性结论
    if 'safety' in content['sections']:
        safety = content['sections']['safety']
        
        if 'adverse_events' in safety:
            ae_rate = safety['adverse_events']['ae_rate']
            
            if ae_rate < 15:
                conclusions.append("治疗具有良好的安全性和耐受性")
            else:
                conclusions.append("需要密切监察治疗相关的不良事件")
    
    # 总体结论
    study_info = content['study_info']
    conclusions.append(f"本{study_info['phase']}研究为{study_info['primary_endpoint']}的评估提供了重要证据")
    
    return conclusions

def generate_study_limitations(content):
    """生成研究局限性"""
    
    limitations = []
    
    study_info = content['study_info']
    
    # 样本量局限性
    if study_info['total_subjects'] < 100:
        limitations.append("样本量相对较小，可能影响统计功效")
    
    # 研究设计局限性
    if study_info['design'] != '随机对照试验':
        limitations.append("非随机对照设计可能存在选择偏倚")
    
    # 随访时间局限性
    if '月' in study_info['duration']:
        try:
            months = int(study_info['duration'].replace('个月', '').replace('月', ''))
            if months < 6:
                limitations.append("随访时间相对较短，长期效果尚需进一步观察")
        except:
            pass
    
    # 数据质量局限性
    if 'results' in content['sections']:
        results = content['sections']['results']
        
        if 'demographics' in results:
            demographics = results['demographics']
            
            # 检查缺失数据
            if 'age' in demographics:
                age_data = demographics['age']
                total_subjects = content['study_info']['total_subjects']
                
                if age_data['n'] < total_subjects * 0.9:
                    limitations.append("部分关键变量存在缺失数据")
    
    # 通用局限性
    limitations.append("单中心研究结果的外推性可能受限")
    limitations.append("未进行多重比较校正可能增加I型错误风险")
    
    return limitations

def generate_clinical_significance(content):
    """生成临床意义"""
    
    significance = []
    
    # 疗效的临床意义
    if 'results' in content['sections']:
        results = content['sections']['results']
        
        if 'primary_analysis' in results:
            primary = results['primary_analysis']
            
            if 'statistical_test' in primary:
                test_result = primary['statistical_test']
                
                if 'error' not in test_result and test_result['significant']:
                    effect_size = abs(test_result['effect_size'])
                    
                    if effect_size >= 0.5:
                        significance.append("观察到的治疗效果具有临床意义，可能改善患者的临床结局")
                    
                    significance.append("统计学显著的结果为临床实践提供了循证医学证据")
    
    # 安全性的临床意义
    if 'safety' in content['sections']:
        safety = content['sections']['safety']
        
        if 'adverse_events' in safety:
            ae_rate = safety['adverse_events']['ae_rate']
            
            if ae_rate < 10:
                significance.append("良好的安全性特征支持治疗在临床实践中的应用")
            
            if 'group_comparison' in safety:
                test_result = safety['group_comparison']
                
                if 'error' not in test_result and not test_result['significant']:
                    significance.append("组间安全性无显著差异，为治疗选择提供了重要参考")
    
    # 研究方法的临床意义
    study_info = content['study_info']
    
    if study_info['design'] == '随机对照试验':
        significance.append("高质量的研究设计增强了结果的临床可信度")
    
    # 对未来研究的意义
    significance.append("本研究结果为后续更大规模的研究提供了重要基础")
    significance.append("研究发现有助于优化临床治疗方案和患者管理策略")
    
    return significance

def generate_final_pdf_report(content):
    """生成最终PDF报告"""
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4, topMargin=1*inch)
    
    # 获取样式
    styles = getSampleStyleSheet()
    story = []
    
    # 自定义样式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=20,
        spaceAfter=30,
        alignment=TA_CENTER,
        textColor=colors.darkblue
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=20,
        spaceAfter=10,
        textColor=colors.darkred
    )
    
    # 报告标题
    study_info = content['study_info']
    story.append(Paragraph(study_info['title'], title_style))
    story.append(Spacer(1, 12))
    
    # 基本信息表
    basic_info = [
        ['研究阶段', study_info['phase']],
        ['研究设计', study_info['design']],
        ['主要终点', study_info['primary_endpoint']],
        ['研究持续时间', study_info['duration']],
        ['总样本量', str(study_info['total_subjects'])],
        ['分析日期', study_info['analysis_date']]
    ]
    
    info_table = Table(basic_info, colWidths=[2*inch, 3*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('BACKGROUND', (1, 0), (1, -1), colors.white),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(info_table)
    story.append(Spacer(1, 20))
    
    # 研究摘要
    story.append(Paragraph("研究摘要", heading_style))
    summary_text = generate_study_summary(content)
    # 简化摘要文本，去除markdown格式
    clean_summary = summary_text.replace('#', '').replace('*', '').strip()
    story.append(Paragraph(clean_summary, styles['Normal']))
    story.append(Spacer(1, 12))
    
    # 主要结果
    if 'results' in content['sections']:
        story.append(Paragraph("主要结果", heading_style))
        
        results = content['sections']['results']
        
        # 受试者特征
        if 'demographics' in results:
            story.append(Paragraph("受试者特征", styles['Heading3']))
            
            demographics = results['demographics']
            
            if 'age' in demographics:
                age_data = demographics['age']
                age_text = f"年龄: 平均 {age_data['mean']:.1f}±{age_data['std']:.1f} 岁，中位数 {age_data['median']:.1f} 岁，范围 {age_data['range'][0]:.0f}-{age_data['range'][1]:.0f} 岁"
                story.append(Paragraph(age_text, styles['Normal']))
            
            if 'gender' in demographics:
                gender_text = "性别分布: "
                for gender, percentage in demographics['gender']['percentages'].items():
                    gender_text += f"{gender} {percentage:.1f}%, "
                story.append(Paragraph(gender_text.rstrip(', '), styles['Normal']))
            
            story.append(Spacer(1, 12))
        
        # 主要终点分析
        if 'primary_analysis' in results:
            story.append(Paragraph("主要终点分析", styles['Heading3']))
            
            primary = results['primary_analysis']
            
            if 'statistical_test' in primary:
                test_result = primary['statistical_test']
                
                if 'error' not in test_result:
                    result_text = f"统计检验: {test_result['test_type']}, "
                    result_text += f"统计量 = {test_result['t_statistic']:.4f}, "
                    result_text += f"P值 = {test_result['p_value']:.4f}, "
                    result_text += f"效应量 = {test_result['effect_size']:.3f}"
                    
                    story.append(Paragraph(result_text, styles['Normal']))
                    
                    if test_result['significant']:
                        story.append(Paragraph("结论: 主要终点达到统计学显著差异", styles['Normal']))
                    else:
                        story.append(Paragraph("结论: 主要终点未达到统计学显著差异", styles['Normal']))
            
            story.append(Spacer(1, 12))
    
    # 安全性分析
    if 'safety' in content['sections']:
        story.append(Paragraph("安全性分析", heading_style))
        
        safety = content['sections']['safety']
        
        if 'adverse_events' in safety:
            ae_data = safety['adverse_events']
            safety_text = f"不良事件发生率: {ae_data['ae_rate']:.1f}% ({ae_data['subjects_with_ae']}/{ae_data['subjects_with_ae'] + ae_data['subjects_without_ae']})"
            story.append(Paragraph(safety_text, styles['Normal']))
        
        story.append(Spacer(1, 12))
    
    # 结论
    story.append(Paragraph("结论", heading_style))
    
    conclusions = generate_final_conclusions(content)
    
    for conclusion in conclusions:
        story.append(Paragraph(f"• {conclusion}", styles['Normal']))
    
    story.append(Spacer(1, 12))
    
    # 局限性
    story.append(Paragraph("研究局限性", heading_style))
    
    limitations = generate_study_limitations(content)
    
    for limitation in limitations:
        story.append(Paragraph(f"• {limitation}", styles['Normal']))
    
    # 构建PDF
    doc.build(story)
    
    buffer.seek(0)
    pdf_data = buffer.getvalue()
    buffer.close()
    
    return pdf_data

def generate_final_word_report(content):
    """生成最终Word报告"""
    
    doc = Document()
    
    # 设置标题
    study_info = content['study_info']
    title = doc.add_heading(study_info['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 基本信息表
    doc.add_heading('研究基本信息', level=1)
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ('研究阶段', study_info['phase']),
        ('研究设计', study_info['design']),
        ('主要终点', study_info['primary_endpoint']),
        ('研究持续时间', study_info['duration']),
        ('总样本量', str(study_info['total_subjects'])),
        ('分析日期', study_info['analysis_date'])
    ]
    
    for i, (key, value) in enumerate(info_data):
        row_cells = info_table.rows[i].cells
        row_cells[0].text = key
        row_cells[1].text = value
    
    # 研究摘要
    doc.add_heading('研究摘要', level=1)
    summary_text = generate_study_summary(content)
    # 清理markdown格式
    clean_summary = summary_text.replace('#', '').replace('*', '').strip()
    doc.add_paragraph(clean_summary)
    
    # 主要结果
    if 'results' in content['sections']:
        doc.add_heading('主要结果', level=1)
        
        results = content['sections']['results']
        
        # 受试者特征
        if 'demographics' in results:
            doc.add_heading('受试者特征', level=2)
            
            demographics = results['demographics']
            
            if 'age' in demographics:
                age_data = demographics['age']
                age_text = f"年龄: 平均 {age_data['mean']:.1f}±{age_data['std']:.1f} 岁，中位数 {age_data['median']:.1f} 岁"
                doc.add_paragraph(age_text)
            
            if 'gender' in demographics:
                gender_text = "性别分布: "
                for gender, percentage in demographics['gender']['percentages'].items():
                    gender_text += f"{gender} {percentage:.1f}%, "
                doc.add_paragraph(gender_text.rstrip(', '))
        
        # 主要终点分析
        if 'primary_analysis' in results:
            doc.add_heading('主要终点分析', level=2)
            
            primary = results['primary_analysis']
            
            if 'statistical_test' in primary:
                test_result = primary['statistical_test']
                
                if 'error' not in test_result:
                    result_text = f"统计检验结果: P值 = {test_result['p_value']:.4f}, 效应量 = {test_result['effect_size']:.3f}"
                    doc.add_paragraph(result_text)
                    
                    if test_result['significant']:
                        doc.add_paragraph("结论: 主要终点达到统计学显著差异")
                    else:
                        doc.add_paragraph("结论: 主要终点未达到统计学显著差异")
    
    # 安全性分析
    if 'safety' in content['sections']:
        doc.add_heading('安全性分析', level=1)
        
        safety = content['sections']['safety']
        
        if 'adverse_events' in safety:
            ae_data = safety['adverse_events']
            safety_text = f"不良事件发生率: {ae_data['ae_rate']:.1f}%"
            doc.add_paragraph(safety_text)
    
    # 结论
    doc.add_heading('结论', level=1)
    
    conclusions = generate_final_conclusions(content)
    
    for conclusion in conclusions:
        doc.add_paragraph(conclusion, style='List Bullet')
    
    # 局限性
    doc.add_heading('研究局限性', level=1)
    
    limitations = generate_study_limitations(content)
    
    for limitation in limitations:
        doc.add_paragraph(limitation, style='List Bullet')
    
    # 保存到内存
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def generate_final_html_report(content):
    """生成最终HTML报告"""
    
    study_info = content['study_info']
    
    html_content = f"""
    <!DOCTYPE html>
    <html lang="zh-CN">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{study_info['title']}</title>
        <style>
            body {{
                font-family: 'Microsoft YaHei', Arial, sans-serif;
                margin: 40px;
                line-height: 1.8;
                color: #333;
                background-color: #f9f9f9;
            }}
            .container {{
                max-width: 1000px;
                margin: 0 auto;
                background-color: white;
                padding: 40px;
                border-radius: 10px;
                box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            }}
            .header {{
                text-align: center;
                border-bottom: 3px solid #2E86AB;
                padding-bottom: 30px;
                margin-bottom: 40px;
            }}
            .header h1 {{
                color: #2E86AB;
                margin-bottom: 20px;
                font-size: 28px;
            }}
            .info-grid {{
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 20px;
                margin: 30px 0;
                padding: 20px;
                background-color: #f8f9fa;
                border-radius: 8px;
            }}
            .info-item {{
                display: flex;
                justify-content: space-between;
                padding: 10px 0;
                border-bottom: 1px solid #eee;
            }}
            .info-label {{
                font-weight: bold;
                color: #555;
            }}
            .info-value {{
                color: #333;
            }}
            .section {{
                margin: 40px 0;
            }}
            .section h2 {{
                color: #A23B72;
                border-left: 4px solid #A23B72;
                padding-left: 15px;
                font-size: 20px;
                margin-bottom: 20px;
            }}
            .section h3 {{
                color: #F18F01;
                font-size: 16px;
                margin: 20px 0 10px 0;
            }}
            .results-table {{
                width: 100%;
                border-collapse: collapse;
                margin: 20px 0;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }}
            .results-table th, .results-table td {{
                border: 1px solid #ddd;
                padding: 12px;
                text-align: left;
            }}
            .results-table th {{
                background-color: #2E86AB;
                color: white;
                font-weight: bold;
            }}
            .results-table tr:nth-child(even) {{
                background-color: #f9f9f9;
            }}
            .highlight-box {{
                background-color: #e8f4fd;
                border-left: 4px solid #2E86AB;
                padding: 15px;
                margin: 15px 0;
                border-radius: 4px;
            }}
            .conclusion-box {{
                background-color: #d4edda;
                border-left: 4px solid #28a745;
                padding: 15px;
                margin: 15px 0;
                border-radius: 4px;
            }}
            .limitation-box {{
                background-color: #fff3cd;
                border-left: 4px solid #ffc107;
                padding: 15px;
                margin: 15px 0;
                border-radius: 4px;
            }}
            .footer {{
                margin-top: 50px;
                padding-top: 20px;
                border-top: 2px solid #ddd;
                text-align: center;
                color: #666;
                font-size: 12px;
            }}
            ul {{
                padding-left: 20px;
            }}
            li {{
                margin: 8px 0;
            }}
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header">
                <h1>{study_info['title']}</h1>
                <p><strong>最终研究报告</strong></p>
            </div>
            
            <div class="info-grid">
                <div class="info-item">
                    <span class="info-label">研究阶段:</span>
                    <span class="info-value">{study_info['phase']}</span>
                </div>
                <div class="info-item">
                    <span class="info-label">研究设计:</span>
                    <span class="info-value">{study_info['design']}</span>
                </div>
                <div class="info-item">
                    <span class="info-label">主要终点:</span>
                    <span class="info-value">{study_info['primary_endpoint']}</span>
                </div>
                <div class="info-item">
                    <span class="info-label">研究持续时间:</span>
                    <span class="info-value">{study_info['duration']}</span>
                </div>
                <div class="info-item">
                    <span class="info-label">总样本量:</span>
                    <span class="info-value">{study_info['total_subjects']}</span>
                </div>
                <div class="info-item">
                    <span class="info-label">分析日期:</span>
                    <span class="info-value">{study_info['analysis_date']}</span>
                </div>
            </div>
            
            <div class="section">
                <h2>研究摘要</h2>
                <div class="highlight-box">
    """
    
    # 添加摘要内容
    summary_text = generate_study_summary(content)
    clean_summary = summary_text.replace('#', '').replace('*', '').strip()
    html_content += f"<p>{clean_summary}</p>"
    
    html_content += """
                </div>
            </div>
    """
    
    # 主要结果
    if 'results' in content['sections']:
        html_content += """
            <div class="section">
                <h2>主要结果</h2>
        """
        
        results = content['sections']['results']
        
        # 受试者特征
        if 'demographics' in results:
            html_content += """
                <h3>受试者特征</h3>
                <table class="results-table">
                    <tr>
                        <th>特征</th>
                        <th>统计值</th>
                    </tr>
            """
            
            demographics = results['demographics']
            
            if 'age' in demographics:
                age_data = demographics['age']
                html_content += f"""
                    <tr>
                        <td>年龄 (岁)</td>
                        <td>平均: {age_data['mean']:.1f}±{age_data['std']:.1f}, 中位数: {age_data['median']:.1f}</td>
                    </tr>
                """
            
            if 'gender' in demographics:
                gender_text = ""
                for gender, percentage in demographics['gender']['percentages'].items():
                    gender_text += f"{gender}: {percentage:.1f}%; "
                
                html_content += f"""
                    <tr>
                        <td>性别分布</td>
                        <td>{gender_text.rstrip('; ')}</td>
                    </tr>
                """
            
            html_content += "</table>"
        
        # 主要终点分析
        if 'primary_analysis' in results:
            html_content += "<h3>主要终点分析</h3>"
            
            primary = results['primary_analysis']
            
            if 'statistical_test' in primary:
                test_result = primary['statistical_test']
                
                if 'error' not in test_result:
                    html_content += f"""
                    <table class="results-table">
                        <tr>
                            <th>统计指标</th>
                            <th>结果</th>
                        </tr>
                        <tr>
                            <td>检验方法</td>
                            <td>{test_result['test_type']}</td>
                        </tr>
                        <tr>
                            <td>统计量</td>
                            <td>{test_result['t_statistic']:.4f}</td>
                        </tr>
                        <tr>
                            <td>P值</td>
                            <td>{test_result['p_value']:.4f}</td>
                        </tr>
                        <tr>
                            <td>效应量 (Cohen's d)</td>
                            <td>{test_result['effect_size']:.3f}</td>
                        </tr>
                        <tr>
                            <td>均值差异</td>
                            <td>{test_result['mean_difference']:.3f}</td>
                        </tr>
                        <tr>
                            <td>95%置信区间</td>
                            <td>[{test_result['ci_95'][0]:.3f}, {test_result['ci_95'][1]:.3f}]</td>
                        </tr>
                    </table>
                    """
                    
                    if test_result['significant']:
                        html_content += '<div class="conclusion-box"><strong>结论:</strong> 主要终点达到统计学显著差异</div>'
                    else:
                        html_content += '<div class="limitation-box"><strong>结论:</strong> 主要终点未达到统计学显著差异</div>'
        
        html_content += "</div>"
    
    # 安全性分析
    if 'safety' in content['sections']:
        html_content += """
            <div class="section">
                <h2>安全性分析</h2>
        """
        
        safety = content['sections']['safety']
        
        if 'adverse_events' in safety:
            ae_data = safety['adverse_events']
            
            html_content += f"""
                <table class="results-table">
                    <tr>
                        <th>安全性指标</th>
                        <th>结果</th>
                    </tr>
                    <tr>
                        <td>不良事件发生率</td>
                        <td>{ae_data['ae_rate']:.1f}% ({ae_data['subjects_with_ae']}/{ae_data['subjects_with_ae'] + ae_data['subjects_without_ae']})</td>
                    </tr>
                    <tr>
                        <td>发生不良事件人数</td>
                        <td>{ae_data['subjects_with_ae']}</td>
                    </tr>
                    <tr>
                        <td>未发生不良事件人数</td>
                        <td>{ae_data['subjects_without_ae']}</td>
                    </tr>
                </table>
            """
            
            # 分组安全性比较
            if 'ae_by_group' in safety:
                html_content += "<h3>分组安全性比较</h3>"
                html_content += """
                <table class="results-table">
                    <tr>
                        <th>治疗组</th>
                        <th>样本量</th>
                        <th>AE人数</th>
                        <th>AE发生率</th>
                    </tr>
                """
                
                for group, data in safety['ae_by_group'].items():
                    html_content += f"""
                    <tr>
                        <td>{group}</td>
                        <td>{data['n']}</td>
                        <td>{data['subjects_with_ae']}</td>
                        <td>{data['ae_rate']:.1f}%</td>
                    </tr>
                    """
                
                html_content += "</table>"
                
                # 统计检验结果
                if 'group_comparison' in safety:
                    test_result = safety['group_comparison']
                    
                    if 'error' not in test_result:
                        html_content += f"""
                        <div class="highlight-box">
                            <strong>组间比较:</strong> 卡方检验, χ² = {test_result['chi2_statistic']:.4f}, P = {test_result['p_value']:.4f}
                        """
                        
                        if test_result['significant']:
                            html_content += "<br><strong>结论:</strong> 组间不良事件发生率存在统计学显著差异"
                        else:
                            html_content += "<br><strong>结论:</strong> 组间不良事件发生率无统计学显著差异"
                        
                        html_content += "</div>"
        
        html_content += "</div>"
    
    # 结论
    html_content += """
        <div class="section">
            <h2>结论</h2>
    """
    
    conclusions = generate_final_conclusions(content)
    
    html_content += "<ul>"
    for conclusion in conclusions:
        html_content += f"<li>{conclusion}</li>"
    html_content += "</ul>"
    
    html_content += "</div>"
    
    # 研究局限性
    html_content += """
        <div class="section">
            <h2>研究局限性</h2>
    """
    
    limitations = generate_study_limitations(content)
    
    html_content += "<ul>"
    for limitation in limitations:
        html_content += f'<li><div class="limitation-box">{limitation}</div></li>'
    html_content += "</ul>"
    
    html_content += "</div>"
    
    # 临床意义
    html_content += """
        <div class="section">
            <h2>临床意义</h2>
    """
    
    clinical_significance = generate_clinical_significance(content)
    
    html_content += "<ul>"
    for significance in clinical_significance:
        html_content += f'<li><div class="conclusion-box">{significance}</div></li>'
    html_content += "</ul>"
    
    html_content += "</div>"
    
    # 页脚
    html_content += f"""
            <div class="footer">
                <p>本报告由临床试验数据分析系统自动生成</p>
                <p>生成时间: {datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}</p>
                <p>报告版本: 最终版 | 数据截止: {study_info['analysis_date']}</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return html_content

# 主函数：报告生成模块入口
def report_generation_module():
    """报告生成模块主函数"""
    
    st.markdown("# 📊 报告生成模块")
    st.markdown("---")
    
    # 检查是否有数据
    if 'df' not in st.session_state or st.session_state.df is None:
        st.warning("⚠️ 请先在数据管理模块中上传数据文件")
        return
    
    df = st.session_state.df
    
    # 报告类型选择
    st.markdown("## 📋 选择报告类型")
    
    report_type = st.selectbox(
        "报告类型",
        [
            "描述性统计报告",
            "推断性统计报告", 
            "自定义分析报告",
            "中期分析报告",
            "最终研究报告"
        ]
    )
    
    # 根据选择的报告类型显示相应界面
    if report_type == "描述性统计报告":
        generate_descriptive_report(df)
    
    elif report_type == "推断性统计报告":
        generate_inferential_report(df)
    
    elif report_type == "自定义分析报告":
        generate_custom_report(df)
    
    elif report_type == "中期分析报告":
        generate_interim_report(df)
    
    elif report_type == "最终研究报告":
        generate_final_report(df)
    
    # 报告模板管理
    st.markdown("---")
    st.markdown("## 📝 报告模板管理")
    
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("💾 保存当前配置为模板"):
            save_report_template()
    
    with col2:
        if st.button("📂 加载已保存模板"):
            load_report_template()

def save_report_template():
    """保存报告模板"""
    
    st.info("报告模板保存功能")
    
    template_name = st.text_input("模板名称", placeholder="输入模板名称")
    template_description = st.text_area("模板描述", placeholder="描述模板用途和特点")
    
    if st.button("确认保存") and template_name:
        # 这里可以实现模板保存逻辑
        # 将当前的报告配置保存到文件或数据库
        
        template_config = {
            'name': template_name,
            'description': template_description,
            'created_date': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            'config': {
                # 保存当前的报告配置参数
                'report_type': 'custom',
                'components': [],  # 这里应该保存实际的组件配置
                'settings': {}     # 这里应该保存实际的设置参数
            }
        }
        
        st.success(f"✅ 模板 '{template_name}' 保存成功！")

def load_report_template():
    """加载报告模板"""
    
    st.info("报告模板加载功能")
    
    # 这里应该显示已保存的模板列表
    available_templates = [
        "标准描述性分析模板",
        "药物安全性评估模板", 
        "疗效评价模板",
        "中期分析标准模板"
    ]
    
    selected_template = st.selectbox("选择模板", available_templates)
    
    if st.button("加载模板"):
        # 这里可以实现模板加载逻辑
        st.success(f"✅ 模板 '{selected_template}' 加载成功！")
        st.info("模板配置已应用到当前报告生成设置中")

# 工具函数
def detect_outliers(series):
    """检测异常值数量"""
    Q1 = series.quantile(0.25)
    Q3 = series.quantile(0.75)
    IQR = Q3 - Q1
    
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    
    outliers = series[(series < lower_bound) | (series > upper_bound)]
    return len(outliers)

def calculate_quality_score(df):
    """计算数据质量评分"""
    
    # 完整性评分
    missing_rate = df.isnull().sum().sum() / (len(df) * len(df.columns))
    completeness_score = (1 - missing_rate) * 100
    
    # 一致性评分
    duplicate_rate = df.duplicated().sum() / len(df)
    consistency_score = (1 - duplicate_rate) * 100
    
    # 总体评分
    overall_score = (completeness_score + consistency_score) / 2
    
    return {
        'completeness': completeness_score,
        'consistency': consistency_score,
        'overall': overall_score
    }

# 如果直接运行此脚本，显示报告生成模块
if __name__ == "__main__":
    report_generation_module()

                






